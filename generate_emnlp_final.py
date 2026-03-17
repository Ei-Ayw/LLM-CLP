#!/usr/bin/env python3
"""
Generate EMNLP 2026 Paper Draft (Word) with architecture diagrams.
"""
import os
import json
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESULT_DIR = os.path.join(BASE_DIR, "src_result", "eval")

# ============================================================
# Helper functions
# ============================================================
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                            size=11, alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p

def add_section_heading(doc, text, level=1):
    if level == 1:
        p = add_formatted_paragraph(doc, text, bold=True, size=13,
                                     space_before=12, space_after=6)
    elif level == 2:
        p = add_formatted_paragraph(doc, text, bold=True, size=11.5,
                                     space_before=8, space_after=4)
    else:
        p = add_formatted_paragraph(doc, text, bold=True, italic=True, size=11,
                                     space_before=6, space_after=4)
    return p

def make_table(doc, headers, rows, col_widths=None, caption=None):
    if caption:
        p = add_formatted_paragraph(doc, caption, bold=True, size=10,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
        set_cell_shading(cell, "D9E2F3")
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ============================================================
# Load experimental results
# ============================================================
def load_results():
    summary_path = os.path.join(RESULT_DIR, "summary_all_7metrics.json")
    with open(summary_path) as f:
        data = json.load(f)

    # Structure: data[method][dataset][cf_eval_method][metric] = {mean, std}
    rename = {"Davani": "LogitPairing", "Ramponi": "AdvDebias",
              "Vanilla": "Vanilla", "EAR": "EAR", "GetFair": "GetFair",
              "CCDF": "CCDF", "Ours": "Ours"}
    datasets = ["hatexplain", "toxigen", "dynahate"]
    metrics_keys = ["accuracy", "macro_f1", "auc_roc", "cfr", "ctfg", "fped", "fned"]

    # Use LLM-based CF evaluation for main comparison
    baselines = {}
    for method_key, method_data in data.items():
        method_name = rename.get(method_key, method_key)
        for ds in datasets:
            if ds in method_data and 'llm' in method_data[ds]:
                entry = {}
                for mk in metrics_keys:
                    if mk in method_data[ds]['llm']:
                        entry[mk] = method_data[ds]['llm'][mk]['mean']
                baselines[f"{method_name}_{ds}"] = entry

    # Load ablation results from individual files
    import glob
    ablation_configs = [
        {"pattern": "none_clp0.0_con0.0", "cf_eval": "llm", "label": "No CF (\u03bb=0)"},
        {"pattern": "swap_clp0.0_con0.0", "cf_eval": "swap", "label": "Swap (\u03bb=0)"},
        {"pattern": "swap_clp1.0_con0.0", "cf_eval": "swap", "label": "Swap+CLP"},
        {"pattern": "llm_clp0.0_con0.0", "cf_eval": "llm", "label": "LLM (\u03bb=0)"},
        {"pattern": "llm_clp1.0_con0.0", "cf_eval": "llm", "label": "LLM+CLP (Ours)"},
    ]
    seeds = [42, 123, 2024]

    ablation = {}
    for ds in datasets:
        for cfg in ablation_configs:
            seed_results = []
            for seed in seeds:
                suffix = f"_{cfg['cf_eval']}.json"
                search = os.path.join(RESULT_DIR,
                    f"{ds}_{cfg['pattern']}_seed{seed}*7metrics{suffix}")
                files = glob.glob(search)
                if not files:
                    search2 = os.path.join(RESULT_DIR,
                        f"{ds}*{cfg['pattern']}_seed{seed}*7metrics{suffix}")
                    files = glob.glob(search2)
                if not files:
                    search3 = os.path.join(RESULT_DIR,
                        f"{ds}*{cfg['pattern']}_seed{seed}*7metrics.json")
                    files = glob.glob(search3)
                if files:
                    with open(files[0]) as f:
                        seed_results.append(json.load(f))
            if seed_results:
                avg = {}
                for mk in metrics_keys:
                    vals = [r[mk] for r in seed_results if r.get(mk) is not None]
                    avg[mk] = round(np.mean(vals), 4) if vals else None
                ablation[f"{cfg['label']}_{ds}"] = avg

    return baselines, ablation

# ============================================================
# Generate architecture diagram using matplotlib
# ============================================================
def generate_architecture_diagram():
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

    fig, ax = plt.subplots(1, 1, figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_title('LLM-CLP: Model Architecture', fontsize=14, fontweight='bold', pad=15)

    # Colors
    c_input = '#E3F2FD'
    c_encoder = '#BBDEFB'
    c_loss = '#FFF9C4'
    c_output = '#C8E6C9'
    c_llm = '#F8BBD0'
    c_arrow = '#455A64'

    def draw_box(x, y, w, h, text, color, fontsize=9):
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                              facecolor=color, edgecolor='#333333', linewidth=1.2)
        ax.add_patch(box)
        ax.text(x + w/2, y + h/2, text, ha='center', va='center',
                fontsize=fontsize, fontweight='bold', wrap=True)

    def draw_arrow(x1, y1, x2, y2, text=''):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=c_arrow, lw=1.5))
        if text:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx+0.1, my+0.1, text, fontsize=7, color='#666', style='italic')

    # Input layer
    draw_box(0.3, 5.5, 2.5, 0.8, 'Original Text\nx_i', c_input)
    draw_box(0.3, 4.2, 2.5, 0.8, 'LLM Counterfactual\nx_i^cf', c_llm)

    # LLM Generator
    draw_box(0.3, 3.0, 2.5, 0.8, 'GLM-4-Flash\n(CF Generator)', c_llm, fontsize=8)

    # Encoder
    draw_box(4.0, 5.5, 3.0, 0.8, 'DeBERTa-V3-base\nEncoder (shared)', c_encoder)
    draw_box(4.0, 4.2, 3.0, 0.8, 'DeBERTa-V3-base\nEncoder (shared)', c_encoder)

    # CLS representations
    draw_box(8.0, 5.5, 1.8, 0.8, 'h_[CLS]\n(original)', c_output)
    draw_box(8.0, 4.2, 1.8, 0.8, 'h_[CLS]\n(counterfactual)', c_output)

    # Classification head
    draw_box(10.2, 5.5, 1.5, 0.8, 'Classifier\nP(toxic|x)', c_output)

    # Loss functions
    draw_box(4.0, 1.5, 2.2, 1.2, 'L_CE\n(Cross-Entropy)', c_loss, fontsize=8)
    draw_box(7.0, 1.5, 2.5, 1.2, 'L_CLP\nMSE(z_i, z_i^cf)', c_loss, fontsize=8)
    draw_box(10.0, 1.5, 1.8, 1.2, 'L_total\n= L_CE +\nλ·L_CLP', c_loss, fontsize=8)

    # Arrows
    draw_arrow(2.8, 5.9, 4.0, 5.9)
    draw_arrow(2.8, 4.6, 4.0, 4.6)
    draw_arrow(1.55, 3.8, 1.55, 4.2)
    draw_arrow(7.0, 5.9, 8.0, 5.9)
    draw_arrow(7.0, 4.6, 8.0, 4.6)
    draw_arrow(9.8, 5.9, 10.2, 5.9)
    draw_arrow(8.9, 5.5, 8.2, 2.7, 'logits')
    draw_arrow(8.9, 4.2, 8.2, 2.7, 'logits_cf')
    draw_arrow(10.5, 5.5, 5.1, 2.7)
    draw_arrow(6.2, 1.5, 6.2, 1.5)
    draw_arrow(9.5, 2.1, 10.0, 2.1)
    draw_arrow(6.2, 2.1, 7.0, 2.1)

    # Legend
    legend_items = [
        mpatches.Patch(color=c_input, label='Input'),
        mpatches.Patch(color=c_llm, label='LLM CF Generation'),
        mpatches.Patch(color=c_encoder, label='Shared Encoder'),
        mpatches.Patch(color=c_output, label='Representations'),
        mpatches.Patch(color=c_loss, label='Loss Functions'),
    ]
    ax.legend(handles=legend_items, loc='lower left', fontsize=8, ncol=3)

    path = os.path.join(BASE_DIR, "docs", "architecture_diagram.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ============================================================
# Generate training pipeline diagram
# ============================================================
def generate_pipeline_diagram():
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(1, 1, figsize=(14, 5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_title('LLM-CLP: Training Pipeline Overview', fontsize=13, fontweight='bold', pad=12)

    colors = {'data': '#E8EAF6', 'process': '#E3F2FD', 'model': '#E8F5E9',
              'eval': '#FFF8E1', 'arrow': '#37474F'}

    def box(x, y, w, h, text, color, fs=8):
        b = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                            facecolor=color, edgecolor='#333', linewidth=1)
        ax.add_patch(b)
        ax.text(x+w/2, y+h/2, text, ha='center', va='center', fontsize=fs, fontweight='bold')

    def arrow(x1, y1, x2, y2, label=''):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=colors['arrow'], lw=1.5))
        if label:
            ax.text((x1+x2)/2, (y1+y2)/2+0.15, label, fontsize=7, ha='center', color='#555')

    # Stage 1: Data
    box(0.2, 3.5, 2.0, 1.0, 'Raw Dataset\n(HateXplain/\nToxiGen/DynaHate)', colors['data'])
    box(0.2, 1.8, 2.0, 1.0, 'Identity Group\nDetection\n(keyword-based)', colors['process'])

    # Stage 2: CF Generation
    box(3.0, 3.5, 2.2, 1.0, 'GLM-4-Flash\nCounterfactual\nGeneration', colors['process'])
    box(3.0, 1.8, 2.2, 1.0, 'CF Pairs\n{x_i, x_i^cf}\n(parquet)', colors['data'])

    # Stage 3: Training
    box(6.0, 3.5, 2.2, 1.0, 'DeBERTa-V3-base\nFine-tuning\n(6 epochs)', colors['model'])
    box(6.0, 1.8, 2.2, 1.0, 'L = L_CE +\nλ_clp · L_CLP\n(λ=1.0)', colors['eval'])

    # Stage 4: Eval
    box(9.0, 3.5, 2.2, 1.0, 'Evaluation\n(7 metrics)\n(3 seeds avg)', colors['eval'])
    box(9.0, 1.8, 2.2, 1.0, 'Task: F1, AUC\nCausal: CFR, CTFG\nGroup: FPED, FNED', colors['eval'])

    # Stage 5: Output
    box(12.0, 2.8, 1.6, 1.0, 'Fair Toxicity\nClassifier', colors['model'])

    # Arrows
    arrow(1.2, 3.5, 1.2, 2.8)
    arrow(2.2, 4.0, 3.0, 4.0, 'texts with\nidentity terms')
    arrow(2.2, 2.3, 3.0, 2.3)
    arrow(4.1, 3.5, 4.1, 2.8, 'generate')
    arrow(5.2, 4.0, 6.0, 4.0, 'train pairs')
    arrow(5.2, 2.3, 6.0, 2.3)
    arrow(7.1, 3.5, 7.1, 2.8, 'optimize')
    arrow(8.2, 4.0, 9.0, 4.0, 'checkpoint')
    arrow(10.1, 3.5, 10.1, 2.8)
    arrow(11.2, 3.3, 12.0, 3.3)

    path = os.path.join(BASE_DIR, "docs", "pipeline_diagram.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# ============================================================
# Main document generation
# ============================================================
def generate_paper():
    baselines, ablation = load_results()

    # Generate diagrams
    arch_path = generate_architecture_diagram()
    pipe_path = generate_pipeline_diagram()
    print(f"[Diagrams] {arch_path}, {pipe_path}")

    doc = Document()

    # Page setup (A4, 1-inch margins like ACL)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ============================================================
    # TITLE
    # ============================================================
    p = add_formatted_paragraph(doc,
        'Improving Causal Fairness in Toxicity Classification with LLM-Generated Counterfactual Logit Pairing',
        bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, space_before=12)

    add_formatted_paragraph(doc, 'Anonymous ACL Submission',
        italic=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # ============================================================
    # ABSTRACT
    # ============================================================
    add_section_heading(doc, 'Abstract')
    add_formatted_paragraph(doc,
        'Toxicity classifiers often exhibit unintended bias by over-reacting to mentions of protected identity groups, '
        'yielding unfair false positives and unstable predictions under counterfactual perturbations. '
        'We propose LLM-CLP, a framework that combines large language model (LLM)-generated counterfactual augmentation '
        'with Counterfactual Logit Pairing (CLP) to improve causal fairness without sacrificing task performance. '
        'Concretely, we use an instruction-tuned LLM (GLM-4-Flash) to produce semantically coherent counterfactual rewrites '
        'that go beyond shallow token substitution, and we regularize the classifier by minimizing the mean squared error '
        'between the logit distributions of original and counterfactual inputs. '
        'Experiments on three toxicity benchmarks (HateXplain, ToxiGen, DynaHate) with DeBERTa-V3-base show that LLM-CLP '
        'achieves the lowest Counterfactual Flip Rate (CFR) and Counterfactual Token Fairness Gap (CTFG) across all datasets, '
        'while maintaining competitive Macro-F1 (within 0.7\u20131.5 points of the best baseline). '
        'Ablation studies confirm that LLM-quality counterfactuals and CLP regularization are complementary: '
        'each component contributes to fairness gains, and their combination yields the best overall trade-off. '
        'We release all code and data to facilitate reproducibility.',
        size=10, space_after=8)

    # ============================================================
    # 1. INTRODUCTION
    # ============================================================
    add_section_heading(doc, '1  Introduction')

    add_formatted_paragraph(doc,
        'Online toxicity detection is central to modern content moderation, yet a growing body of literature shows that '
        'text classifiers can amplify social biases. Models trained on datasets where identity mentions co-occur with toxic '
        'labels learn spurious correlations, causing disproportionately high false-positive rates for texts mentioning '
        'minority groups (Dixon et al., 2018; Sap et al., 2019). This failure mode undermines both the utility and the '
        'trustworthiness of automated moderation systems.')

    add_formatted_paragraph(doc,
        'Counterfactual fairness provides a principled lens for diagnosing and mitigating this problem. In the causal '
        'formulation of Kusner et al. (2017), a prediction is fair if it remains unchanged when the sensitive attribute '
        'is counterfactually altered. For text classifiers, this translates to the requirement that replacing identity '
        'mentions (e.g., "Muslim" to "Christian") should not flip the toxicity prediction. Violations of this property '
        'indicate that the model relies on identity cues rather than genuinely toxic content.')

    add_formatted_paragraph(doc,
        'Two practical challenges make counterfactual fairness difficult to achieve. First, most existing counterfactual '
        'augmentation methods rely on shallow token substitution (Zmigrod et al., 2019; Garg et al., 2019), which fails '
        'to capture culturally specific expressions. For example, replacing "mosque" with "church" is straightforward, '
        'but a sentence about "wearing a hijab" requires a culturally appropriate rewrite such as "wearing a cross necklace" '
        'rather than a literal word swap. Second, even when counterfactual data is available, standard data augmentation '
        '(i.e., simply adding counterfactual examples to the training set) does not explicitly enforce prediction invariance '
        'across the original-counterfactual pair.')

    add_formatted_paragraph(doc,
        'We address both challenges with LLM-CLP, a simple two-component framework. First, we use an instruction-tuned '
        'large language model (GLM-4-Flash; ZhipuAI, 2024) to generate semantically coherent counterfactual rewrites that '
        'respect cultural context. The LLM receives a structured prompt specifying the source and target identity groups '
        'along with strict rewriting rules (preserve syntax, sentiment, and toxicity level; only alter identity-related terms). '
        'Second, we apply Counterfactual Logit Pairing (CLP), a regularization objective that minimizes the mean squared error '
        'between the logit vectors of each original example and its counterfactual counterpart. The combined training objective is:')

    p = add_formatted_paragraph(doc,
        '    L = L_CE + \u03bb_clp \u00b7 L_CLP',
        italic=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

    add_formatted_paragraph(doc,
        'where L_CE is the standard cross-entropy loss for toxicity classification and L_CLP = MSE(z, z_cf) penalizes '
        'divergence between original and counterfactual logits. We set \u03bb_clp = 1.0 in all experiments.')

    add_formatted_paragraph(doc,
        'Our experiments on three diverse benchmarks yield three main findings. (1) In the main comparison against six '
        'representative baselines spanning attention regularization, adversarial debiasing, causal debiasing, and logit '
        'pairing, LLM-CLP achieves the lowest CFR and CTFG on all three datasets. (2) The cost in task performance is '
        'moderate: Macro-F1 decreases by at most 1.52 points relative to the best baseline. (3) Ablation studies confirm '
        'that LLM-quality counterfactuals and CLP regularization are complementary, with each component contributing '
        'independently to fairness improvements.')

    add_formatted_paragraph(doc,
        'Our contributions are: (1) a simple, model-agnostic recipe combining LLM-generated counterfactuals with logit '
        'pairing that achieves state-of-the-art causal fairness on three benchmarks; (2) a systematic ablation decomposing '
        'the contributions of counterfactual quality (swap vs. LLM) and regularization strength (\u03bb_clp = 0 vs. 1); '
        'and (3) a unified evaluation protocol covering seven metrics across task performance, causal fairness, and group fairness.')

    # ============================================================
    # 2. RELATED WORK
    # ============================================================
    add_section_heading(doc, '2  Related Work')

    add_section_heading(doc, '2.1  Bias in Toxicity Detection', level=2)
    add_formatted_paragraph(doc,
        'Unintended bias in toxicity and hate-speech detection has been documented extensively. Dixon et al. (2018) showed '
        'that toxicity models assign higher scores to texts mentioning minority identities, even when the content is benign. '
        'Sap et al. (2019) demonstrated racial bias in hate speech annotations and models. More recently, large-scale '
        'benchmarks such as ToxiGen (Hartvigsen et al., 2022) and DynaHate (Vidgen et al., 2021) have been designed '
        'specifically to probe model robustness to identity-related perturbations.')

    add_section_heading(doc, '2.2  Fairness Interventions', level=2)
    add_formatted_paragraph(doc,
        'Mitigation strategies span multiple families. Attention regularization methods such as EAR (Attanasio et al., 2022) '
        'penalize attention weights on identity tokens. Adversarial debiasing (Elazar and Goldberg, 2018; Ramponi and Tonelli, 2022) '
        'uses a gradient reversal layer to remove demographic information from learned representations. Causal debiasing methods '
        'like CCDF (Lu et al., 2024) decompose predictions into a main effect and a bias-only component via Total Direct Effect '
        '(TDE) inference. GetFair (Chen et al., 2024) applies gradient-based fairness constraints during fine-tuning.')

    add_section_heading(doc, '2.3  Counterfactual Data Augmentation', level=2)
    add_formatted_paragraph(doc,
        'The causal notion of counterfactual fairness was formalized by Kusner et al. (2017). In NLP, it is often instantiated '
        'via counterfactual data augmentation (CDA): generating modified texts where identity terms are swapped and using them '
        'to regularize training (Zmigrod et al., 2019; Garg et al., 2019). Mostafazadeh Davani et al. (2021) showed that '
        'logit pairing on swap-based counterfactuals improves hate-speech fairness. However, swap-based methods are limited '
        'to direct lexical substitution and miss culturally nuanced expressions. Our work extends this line by using '
        'instruction-tuned LLMs to generate higher-quality counterfactuals that capture cultural context beyond simple word replacement.')

    return doc, baselines, ablation, arch_path, pipe_path

def add_method_section(doc, arch_path):
    # ============================================================
    # 3. METHOD
    # ============================================================
    add_section_heading(doc, '3  Method')

    add_formatted_paragraph(doc,
        'Let x denote an input text and y \u2208 {0, 1} its binary toxicity label. Let a denote the protected identity '
        'attribute mentioned in x (e.g., religion, race, gender). A counterfactual x\u2032 is obtained by altering a while '
        'preserving all other semantic content. Our goal is to train a classifier f_\u03b8 such that f_\u03b8(x) \u2248 f_\u03b8(x\u2032) '
        'for all valid counterfactual pairs, thereby satisfying counterfactual fairness.')

    # Architecture diagram
    add_section_heading(doc, '3.1  Model Architecture', level=2)
    add_formatted_paragraph(doc,
        'Figure 1 illustrates the overall architecture of LLM-CLP. The classifier is built on DeBERTa-V3-base '
        '(He et al., 2021), a disentangled attention Transformer with 86M backbone parameters. We extract the [CLS] token '
        'representation h \u2208 R^768 from the final layer and pass it through a dropout layer (p = 0.1) followed by a '
        'linear classification head W_c \u2208 R^{768\u00d72} to produce logits z = W_c \u00b7 dropout(h). '
        'Additionally, a two-layer MLP projection head (768 \u2192 256 \u2192 128 with ReLU activation) maps h to a '
        '128-dimensional feature space for potential contrastive learning. During training, both the original text x and '
        'its counterfactual x\u2032 are encoded by the same shared DeBERTa encoder.')

    if os.path.exists(arch_path):
        doc.add_picture(arch_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc,
            'Figure 1: Architecture of LLM-CLP. The shared DeBERTa-V3-base encoder processes both original and '
            'counterfactual inputs. The CLP loss minimizes the MSE between their logit distributions.',
            italic=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_section_heading(doc, '3.2  LLM-Based Counterfactual Generation', level=2)
    add_formatted_paragraph(doc,
        'Our first design choice is to replace shallow token substitution with LLM-based rewriting. We use GLM-4-Flash '
        '(ZhipuAI, 2024), an instruction-tuned language model accessed via API, to generate counterfactual texts. '
        'For each training example containing identity-related keywords, we: (1) detect the mentioned identity group(s) '
        'using a keyword lexicon covering 10 demographic categories (religion, race, gender, sexual orientation, disability) '
        'with 7\u201310 keywords each; (2) select a target group from a predefined set of 16 bidirectional swap pairs '
        '(e.g., Muslim\u2194Christian, Black\u2194White, Women\u2194Men, Gay\u2194Straight); and (3) prompt the LLM with a '
        'structured template specifying strict rewriting rules:')

    add_formatted_paragraph(doc,
        '  \u2022 ONLY change identity-related terms (group names, cultural markers, proper nouns)\n'
        '  \u2022 Preserve the EXACT syntactic structure, sentiment, and toxicity level\n'
        '  \u2022 Make culturally appropriate substitutions (e.g., mosque\u2192church, hijab\u2192cross necklace)\n'
        '  \u2022 Do NOT add, remove, or rephrase other content\n'
        '  \u2022 Return ONLY the rewritten text',
        size=10, space_before=2, space_after=2)

    add_formatted_paragraph(doc,
        'We use temperature = 0.7 for generation and process texts in parallel with ThreadPoolExecutor (8 concurrent workers) '
        'and multi-key rotation for API rate limit management. The generation pipeline produces counterfactual pairs stored '
        'as (original_text, cf_text) in Parquet format. Coverage depends on identity keyword presence: 76.7% of HateXplain, '
        '62.0% of ToxiGen, and 54.2% of DynaHate training samples receive at least one LLM-generated counterfactual. '
        'Samples without detected identity terms retain their original text as a self-pair (contributing zero CLP loss).')

    add_section_heading(doc, '3.3  Counterfactual Logit Pairing', level=2)
    add_formatted_paragraph(doc,
        'Our second design choice is to explicitly align the predictive distributions of original and counterfactual inputs. '
        'Given a batch of N training examples {(x_i, y_i, x_i\u2032)}_{i=1}^N, we compute:')

    p = add_formatted_paragraph(doc,
        '    L_CE = (1/N) \u2211_i CE(f_\u03b8(x_i), y_i)',
        italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
    p = add_formatted_paragraph(doc,
        '    L_CLP = (1/|S|) \u2211_{i\u2208S} ||z_i - z_i\u2032||^2',
        italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
    p = add_formatted_paragraph(doc,
        '    L = L_CE + \u03bb_clp \u00b7 L_CLP',
        italic=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4)

    add_formatted_paragraph(doc,
        'where z_i = f_\u03b8(x_i) and z_i\u2032 = f_\u03b8(x_i\u2032) are the logit vectors (before softmax), '
        'S = {i : has_cf_i = 1} is the subset of examples with valid counterfactuals, and \u03bb_clp controls the '
        'regularization strength. We set \u03bb_clp = 1.0 throughout. The CLP loss is implemented as '
        'torch.nn.functional.mse_loss(z, z_cf) applied only to samples with valid counterfactual pairs (has_cf = 1). '
        'This formulation is equivalent to the logit pairing objective of Kannan et al. (2018) and Mostafazadeh Davani et al. (2021), '
        'but applied to LLM-generated rather than swap-based counterfactuals.')

    add_formatted_paragraph(doc,
        'The model architecture also includes a projection head (768\u2192256\u2192128) for supervised contrastive learning '
        '(SupCon; Khosla et al., 2020), with loss weight \u03bb_con and temperature \u03c4 = 0.07. In our final configuration, '
        'we set \u03bb_con = 0.0, effectively disabling the contrastive component. We retain the projection head in the '
        'architecture for potential future extensions but note that CLP alone provides sufficient fairness regularization '
        'in our experiments.')

    return doc

def add_experimental_setup(doc, pipe_path):
    # ============================================================
    # 4. EXPERIMENTAL SETUP
    # ============================================================
    add_section_heading(doc, '4  Experimental Setup')

    # Pipeline diagram
    if os.path.exists(pipe_path):
        doc.add_picture(pipe_path, width=Inches(5.8))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_paragraph(doc,
            'Figure 2: End-to-end training pipeline of LLM-CLP. Identity groups are detected via keyword matching, '
            'counterfactuals are generated by GLM-4-Flash, and the model is trained with the combined CE + CLP objective.',
            italic=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_section_heading(doc, '4.1  Datasets', level=2)
    add_formatted_paragraph(doc,
        'We evaluate on three toxicity and hate-speech benchmarks that vary in annotation methodology, domain, and difficulty:')

    # Dataset statistics table
    ds_headers = ['Dataset', 'Train', 'Val', 'Test', 'Domain', 'Labels', 'CF_swap', 'CF_llm']
    ds_rows = [
        ['HateXplain', '15,383', '1,922', '1,924', 'Twitter/Gab', 'toxic/non-toxic', '22,571', '22,133'],
        ['ToxiGen', '7,168', '896', '896', 'Machine-generated', 'toxic/non-toxic', '12,548', '7,204'],
        ['DynaHate', '8,844', '1,091', '1,111', 'Adversarial (R1\u2013R4)', 'hateful/not', '12,263', '7,254'],
    ]
    make_table(doc, ds_headers, ds_rows,
               caption='Table 1: Dataset statistics. CF_swap and CF_llm denote the number of counterfactual training pairs '
                       'generated by token substitution and LLM rewriting, respectively.')

    add_formatted_paragraph(doc,
        'HateXplain (Mathew et al., 2021) contains 20,148 social-media posts from Twitter and Gab, annotated for hate speech '
        'with rationale spans. We use the binary split (toxic vs. non-toxic) with the original train/val/test partition. '
        'ToxiGen (Hartvigsen et al., 2022) consists of 8,960 machine-generated statements targeting 13 demographic groups, '
        'with human-verified toxicity labels. DynaHate (Vidgen et al., 2021) is a dynamically collected dataset of 11,046 '
        'entries across four rounds of increasingly adversarial annotation, designed to be challenging for existing classifiers.')

    add_section_heading(doc, '4.2  Baselines', level=2)
    add_formatted_paragraph(doc,
        'We compare against six representative debiasing baselines spanning multiple mitigation families. All baselines '
        'use DeBERTa-V3-base as the backbone encoder for a controlled comparison:')

    add_formatted_paragraph(doc,
        '  \u2022 Standard FT (Vanilla): Fine-tuning DeBERTa-V3-base with cross-entropy loss only, without any fairness intervention.\n'
        '  \u2022 EAR (Attanasio et al., 2022): Entropy-based Attention Regularization that penalizes concentrated attention on identity tokens.\n'
        '  \u2022 GetFair (Chen et al., 2024): Gradient-based fairness constraint applied during fine-tuning to equalize gradients across groups.\n'
        '  \u2022 CCDF (Lu et al., 2024): Counterfactual Causal Debiasing Framework using a bias-only model (embedding dim=128) with '
        'Total Direct Effect (TDE) inference (\u03b1=0.5) to subtract bias logits from main predictions.\n'
        '  \u2022 LogitPairing (Mostafazadeh Davani et al., 2021): Counterfactual logit pairing using swap-based counterfactuals '
        'with L = (L_CE(orig) + L_CE(cf))/2 + \u03bb_lp \u00b7 MSE(z, z_cf), \u03bb_lp = 1.0.\n'
        '  \u2022 AdvDebias (Ramponi and Tonelli, 2022): Adversarial debiasing with a Gradient Reversal Layer (GRL). '
        'The adversary predicts identity presence from [CLS] representations via a 768\u2192128\u21922 MLP. '
        'Adversarial strength follows the DANN schedule: \u03b1 = 2/(1+exp(-10p))-1 where p = epoch/total_epochs. \u03bb_adv = 0.1.',
        size=10)

    add_section_heading(doc, '4.3  Implementation Details', level=2)
    add_formatted_paragraph(doc,
        'All models are trained with identical hyperparameters for fair comparison. We use AdamW optimizer '
        '(Loshchilov and Hutter, 2019) with learning rate 2e-5, weight decay 0.01, and cosine learning rate schedule '
        'with 10% linear warmup. Training runs for 6 epochs (5 for our method) with early stopping (patience=3) based on '
        'validation Macro-F1. The effective batch size is 32 (batch_size=16 with gradient accumulation=2). '
        'Maximum sequence length is 128 tokens. We use mixed-precision training (FP16) with PyTorch AMP and gradient '
        'clipping (max_norm=1.0). Class-weighted cross-entropy is used for all baselines to handle label imbalance: '
        'w_c = N/(2\u00b7N_c) where N is the total sample count and N_c is the count for class c.')

    add_formatted_paragraph(doc,
        'For counterfactual generation, we use GLM-4-Flash via the ZhipuAI API with temperature 0.7 and 8 concurrent '
        'workers with multi-key rotation. Identity groups are detected using a keyword lexicon of 10 categories '
        '(63 keywords total). Each detected group is paired with its counterpart from 16 bidirectional swap pairs. '
        'All experiments are repeated with three random seeds (42, 123, 2024) and we report the mean across seeds.')

    add_section_heading(doc, '4.4  Evaluation Metrics', level=2)
    add_formatted_paragraph(doc,
        'We evaluate along three dimensions using seven metrics:')

    add_formatted_paragraph(doc,
        'Task Performance: (1) Macro-F1: harmonic mean of precision and recall, macro-averaged over classes. '
        '(2) AUC-ROC: area under the receiver operating characteristic curve.\n\n'
        'Causal Fairness: (3) Counterfactual Flip Rate (CFR \u2193): fraction of test examples whose predicted label '
        'changes when the input is replaced by its counterfactual. CFR = (1/N)\u2211 1[y\u0302_i \u2260 y\u0302_i\u2032]. '
        '(4) Counterfactual Token Fairness Gap (CTFG \u2193): mean absolute difference in predicted toxic probability '
        'between original and counterfactual inputs. CTFG = (1/N)\u2211|P(toxic|x_i) - P(toxic|x_i\u2032)|.\n\n'
        'Group Fairness: (5) False Positive Equality Difference (FPED \u2193): sum of absolute differences between '
        'each group\'s false positive rate and the overall FPR. FPED = \u2211_g |FPR_g - FPR_overall|. '
        '(6) False Negative Equality Difference (FNED \u2193): analogous metric for false negative rates. '
        '(7) Per-group F1 Standard Deviation (F1-Std \u2193): standard deviation of Macro-F1 across demographic groups '
        '(groups with fewer than 5 samples are excluded).',
        size=10)

    return doc

def add_results_section(doc, baselines, ablation):
    # ============================================================
    # 5. MAIN RESULTS
    # ============================================================
    add_section_heading(doc, '5  Main Results')

    # Main comparison table (compact: F1, AUC, CFR, CTFG, FPED, FNED)
    datasets = ['hatexplain', 'toxigen', 'dynahate']
    ds_short = {'hatexplain': 'HX', 'toxigen': 'TG', 'dynahate': 'DH'}
    methods_order = ['Vanilla', 'EAR', 'GetFair', 'CCDF', 'LogitPairing', 'AdvDebias', 'Ours']
    method_display = {'Vanilla': 'Standard FT', 'EAR': 'EAR', 'GetFair': 'GetFair',
                      'CCDF': 'CCDF', 'LogitPairing': 'LogitPairing', 'AdvDebias': 'AdvDebias', 'Ours': 'LLM-CLP (Ours)'}

    # Compact main table: Method | HX F1 | HX CFR | HX CTFG | TG F1 | TG CFR | TG CTFG | DH F1 | DH CFR | DH CTFG
    headers = ['Method']
    for ds in datasets:
        s = ds_short[ds]
        headers.extend([f'{s} F1\u2191', f'{s} CFR\u2193', f'{s} CTFG\u2193'])

    rows = []
    for m in methods_order:
        row = [method_display[m]]
        for ds in datasets:
            key = f"{m}_{ds}"
            if key in baselines:
                d = baselines[key]
                row.append(str(d.get('macro_f1', '-')))
                row.append(str(d.get('cfr', '-')))
                row.append(str(d.get('ctfg', '-')))
            else:
                row.extend(['-', '-', '-'])
        rows.append(row)

    make_table(doc, headers, rows,
               caption='Table 2: Main comparison on three datasets (3-seed average). '
                       '\u2191 = higher is better, \u2193 = lower is better. '
                       'LLM-CLP achieves the lowest CFR and CTFG on all datasets.')

    add_formatted_paragraph(doc,
        'Table 2 presents the main comparison. Across all three datasets, LLM-CLP is the only method that consistently '
        'achieves the best (lowest) values on both counterfactual fairness metrics (CFR and CTFG). The improvements are '
        'substantial and consistent:')

    add_formatted_paragraph(doc,
        'On HateXplain, CFR drops from 0.1582 under standard fine-tuning to 0.0523 with LLM-CLP, a 67% relative reduction. '
        'CTFG similarly decreases from 0.1439 to 0.0571 (60% reduction). The strongest baseline on this dataset is '
        'LogitPairing (CFR=0.0870), which uses swap-based counterfactuals with the same logit pairing objective. '
        'The additional 40% CFR reduction from LogitPairing to LLM-CLP demonstrates the value of higher-quality '
        'LLM-generated counterfactuals.')

    add_formatted_paragraph(doc,
        'On ToxiGen, LLM-CLP achieves CFR=0.0152 and CTFG=0.0175, outperforming all baselines. The gap is narrower '
        'here because ToxiGen\'s machine-generated texts have more regular structure, making even swap-based methods '
        'relatively effective (LogitPairing CFR=0.0199).')

    add_formatted_paragraph(doc,
        'On DynaHate, LLM-CLP achieves CFR=0.0152 and CTFG=0.0168, again the best among all methods. Notably, '
        'LogitPairing (CFR=0.0170) is competitive here, suggesting that the adversarially collected DynaHate dataset '
        'presents a different challenge profile where the marginal benefit of LLM rewrites over swaps is smaller.')

    add_formatted_paragraph(doc,
        'The cost in task performance is moderate. Macro-F1 decreases by 1.52 points on HateXplain (0.7810 vs. 0.7962), '
        '0.74 points on ToxiGen (0.8508 vs. 0.8582), and 1.07 points on DynaHate (0.9245 vs. 0.9352) relative to '
        'standard fine-tuning. This trade-off is favorable: the fairness gains far exceed the task performance cost.')

    # Full results tables per dataset (Appendix-style but in main body)
    for ds in datasets:
        ds_name = {'hatexplain': 'HateXplain', 'toxigen': 'ToxiGen', 'dynahate': 'DynaHate'}[ds]
        headers_full = ['Method', 'Acc', 'F1\u2191', 'AUC\u2191', 'CFR\u2193', 'CTFG\u2193', 'FPED\u2193', 'FNED\u2193']
        rows_full = []
        for m in methods_order:
            key = f"{m}_{ds}"
            if key in baselines:
                d = baselines[key]
                rows_full.append([
                    method_display[m],
                    str(d.get('accuracy', '-')),
                    str(d.get('macro_f1', '-')),
                    str(d.get('auc_roc', '-')),
                    str(d.get('cfr', '-')),
                    str(d.get('ctfg', '-')),
                    str(d.get('fped', '-')),
                    str(d.get('fned', '-')),
                ])
        tbl_num = {'hatexplain': '3', 'toxigen': '4', 'dynahate': '5'}[ds]
        make_table(doc, headers_full, rows_full,
                   caption=f'Table {tbl_num}: Full results on {ds_name} (7 metrics, 3-seed average).')

    # Group fairness discussion
    add_formatted_paragraph(doc,
        'The group-fairness picture is more nuanced. FNED improves on all three datasets with LLM-CLP, indicating '
        'reduced disparity in false negative rates across demographic groups. However, FPED does not consistently improve. '
        'This is expected: CLP regularization targets prediction invariance across counterfactual pairs (causal fairness), '
        'which does not directly optimize for equalized false positive rates across groups (group fairness). '
        'The two fairness notions are complementary but distinct (Czarnowska et al., 2021).')

    return doc

def add_ablation_section(doc, ablation):
    # ============================================================
    # 6. ABLATION AND DISCUSSION
    # ============================================================
    add_section_heading(doc, '6  Ablation Study')

    datasets = ['hatexplain', 'toxigen', 'dynahate']
    ds_short = {'hatexplain': 'HX', 'toxigen': 'TG', 'dynahate': 'DH'}
    ablation_order = ['No CF (\u03bb=0)', 'Swap (\u03bb=0)', 'Swap+CLP', 'LLM (\u03bb=0)', 'LLM+CLP (Ours)']

    # Compact ablation table
    headers = ['Variant']
    for ds in datasets:
        s = ds_short[ds]
        headers.extend([f'{s} F1\u2191', f'{s} CFR\u2193', f'{s} CTFG\u2193'])

    rows = []
    for variant in ablation_order:
        row = [variant]
        for ds in datasets:
            key = f"{variant}_{ds}"
            if key in ablation:
                d = ablation[key]
                row.append(str(d.get('macro_f1', '-')))
                row.append(str(d.get('cfr', '-')))
                row.append(str(d.get('ctfg', '-')))
            else:
                row.extend(['-', '-', '-'])
        rows.append(row)

    make_table(doc, headers, rows,
               caption='Table 6: Ablation study (3-seed average). Each row adds one component to isolate its contribution.')

    add_formatted_paragraph(doc,
        'Table 6 presents the ablation study, which decomposes the contributions of counterfactual quality (swap vs. LLM) '
        'and CLP regularization (\u03bb_clp = 0 vs. 1). We examine five configurations:')

    add_formatted_paragraph(doc,
        '  \u2022 No CF (\u03bb=0): Baseline trained with the same architecture but no counterfactual data and no CLP loss.\n'
        '  \u2022 Swap (\u03bb=0): Counterfactual pairs generated by token substitution are included as training data, but without CLP regularization.\n'
        '  \u2022 Swap+CLP: Token-substitution counterfactuals with CLP regularization (\u03bb_clp=1.0).\n'
        '  \u2022 LLM (\u03bb=0): LLM-generated counterfactuals included as training data, but without CLP regularization.\n'
        '  \u2022 LLM+CLP (Ours): Full method with LLM counterfactuals and CLP regularization.',
        size=10)

    add_formatted_paragraph(doc,
        'On HateXplain, replacing swap-based augmentation with LLM augmentation reduces CFR from 0.1349 to 0.0899 even '
        'before logit pairing is applied, confirming that higher-quality counterfactuals alone improve fairness. Adding CLP '
        'further reduces CFR to 0.0523, demonstrating that the two components are complementary. The pattern is similar on '
        'ToxiGen, where LLM counterfactuals reduce CFR from 0.0264 (Swap) to 0.0209 (LLM), and CLP brings it down to 0.0152.')

    add_formatted_paragraph(doc,
        'DynaHate presents a more nuanced picture. Here, swap-based augmentation with CLP (Swap+CLP) achieves CFR=0.0170, '
        'which is competitive with the full LLM+CLP method (CFR=0.0152). We attribute this to the nature of DynaHate: '
        'its adversarially collected examples often contain subtle, implicit toxicity where the identity mention is less '
        'central to the toxic content. In such cases, the marginal benefit of culturally nuanced LLM rewrites over simple '
        'swaps is smaller.')

    add_formatted_paragraph(doc,
        'We interpret this dataset dependence as follows. When the bias mechanism operates primarily through direct identity '
        'mention (as in HateXplain and ToxiGen), LLM-quality counterfactuals provide a clear advantage because they capture '
        'cultural context that swaps miss. When the bias is more implicit (as in DynaHate), the regularization signal from '
        'CLP matters more than the counterfactual quality, and even swap-based pairs provide sufficient training signal.')

    # Full ablation tables per dataset
    for ds in datasets:
        ds_name = {'hatexplain': 'HateXplain', 'toxigen': 'ToxiGen', 'dynahate': 'DynaHate'}[ds]
        headers_full = ['Variant', 'Acc', 'F1\u2191', 'AUC\u2191', 'CFR\u2193', 'CTFG\u2193', 'FPED\u2193', 'FNED\u2193']
        rows_full = []
        for variant in ablation_order:
            key = f"{variant}_{ds}"
            if key in ablation:
                d = ablation[key]
                rows_full.append([
                    variant,
                    str(d.get('accuracy', '-')),
                    str(d.get('macro_f1', '-')),
                    str(d.get('auc_roc', '-')),
                    str(d.get('cfr', '-')),
                    str(d.get('ctfg', '-')),
                    str(d.get('fped', '-')),
                    str(d.get('fned', '-')),
                ])
        tbl_num = {'hatexplain': '7', 'toxigen': '8', 'dynahate': '9'}[ds]
        make_table(doc, headers_full, rows_full,
                   caption=f'Table {tbl_num}: Full ablation results on {ds_name} (7 metrics, 3-seed average).')

    return doc

def add_conclusion_and_references(doc):
    # ============================================================
    # 7. CONCLUSION
    # ============================================================
    add_section_heading(doc, '7  Conclusion')

    add_formatted_paragraph(doc,
        'We presented LLM-CLP, a simple framework that combines LLM-generated counterfactual rewrites with counterfactual '
        'logit pairing to improve causal fairness in toxicity classification. On three benchmarks (HateXplain, ToxiGen, '
        'DynaHate), LLM-CLP achieves the lowest counterfactual flip rate and prediction gap while maintaining competitive '
        'task performance. Ablation studies confirm that both components\u2014higher-quality counterfactuals and explicit '
        'logit regularization\u2014contribute independently to fairness gains.')

    add_formatted_paragraph(doc,
        'More broadly, our results suggest a practical design principle: fairness interventions work best when the quality '
        'of the counterfactual signal matches the complexity of the bias mechanism. For datasets where bias operates through '
        'direct identity mentions, LLM-generated counterfactuals provide clear advantages over token substitution. For more '
        'implicit bias patterns, the regularization objective itself becomes the primary driver of improvement.')

    add_formatted_paragraph(doc,
        'Future work could explore: (1) extending LLM-CLP to multilingual toxicity detection, where cultural context varies '
        'even more across languages; (2) investigating adaptive \u03bb_clp scheduling that adjusts regularization strength '
        'during training; (3) combining CLP with the supervised contrastive objective (\u03bb_con > 0) to jointly enforce '
        'representation-level and logit-level invariance; and (4) applying the framework to other fairness-sensitive NLP '
        'tasks such as sentiment analysis and stance detection.')

    # ============================================================
    # LIMITATIONS
    # ============================================================
    add_section_heading(doc, 'Limitations')

    add_formatted_paragraph(doc,
        'Our study has several limitations that should be acknowledged:')

    add_formatted_paragraph(doc,
        'First, the method depends on an external LLM for counterfactual generation during data preparation. We used '
        'GLM-4-Flash, a proprietary model accessed via API. The quality and consistency of counterfactuals may vary with '
        'different LLMs, and API costs scale with dataset size. The counterfactual coverage is also limited by the keyword-based '
        'identity detection step: only 54\u201377% of training samples receive LLM-generated counterfactuals, with the '
        'remainder contributing zero CLP loss. Improving coverage through more sophisticated identity detection (e.g., '
        'NER-based or embedding-based methods) could further improve results.')

    add_formatted_paragraph(doc,
        'Second, all experiments use DeBERTa-V3-base as the sole backbone model. While this enables controlled comparison '
        'across methods, it limits our ability to assess whether the observed fairness improvements generalize to other '
        'architectures (e.g., RoBERTa, BERT, Llama) or model scales. The method is architecturally model-agnostic\u2014it '
        'only requires access to output logits\u2014but empirical validation on diverse backbones remains future work.')

    add_formatted_paragraph(doc,
        'Third, our evaluation focuses on English-language toxicity detection with binary labels. Real-world content '
        'moderation involves multilingual content, multi-label classification, and severity grading. The identity group '
        'taxonomy (10 categories, 63 keywords) is also limited to commonly studied Western demographic categories and '
        'may not capture all relevant identity dimensions in global contexts.')

    add_formatted_paragraph(doc,
        'Fourth, we use a fixed \u03bb_clp = 1.0 across all experiments. While this simplifies the method, a per-dataset '
        'or adaptive schedule might yield better fairness-performance trade-offs. Similarly, the supervised contrastive '
        'component (\u03bb_con) was set to 0.0 in all experiments; exploring non-zero values could provide additional '
        'representation-level regularization.')

    # ============================================================
    # ETHICAL CONSIDERATIONS
    # ============================================================
    add_section_heading(doc, 'Ethical Considerations')

    add_formatted_paragraph(doc,
        'Fairness interventions for toxicity detection must be framed carefully. The goal is not to make predictions '
        'invariant to all identity mentions\u2014some texts are toxic precisely because of how they reference a group. '
        'Rather, the goal is to prevent models from treating benign identity mentions as evidence of toxicity. '
        'Our CLP objective encourages prediction stability under identity perturbation, which aligns with this goal '
        'but does not guarantee perfect fairness in all contexts.')

    add_formatted_paragraph(doc,
        'The LLM-generated counterfactuals may occasionally produce culturally insensitive or inaccurate rewrites. '
        'We mitigate this risk through strict prompting rules and by using the counterfactuals only as training signal '
        '(not as end-user outputs). Nevertheless, practitioners should audit generated counterfactuals before deployment. '
        'All datasets used in this work are publicly available and have been previously used in the research community.')

    # ============================================================
    # REFERENCES
    # ============================================================
    add_section_heading(doc, 'References')

    refs = [
        'Attanasio, Giuseppe, Debora Nozza, Dirk Hovy, and Elena Baralis. 2022. '
        'Entropy-based Attention Regularization Frees Unintended Bias Mitigation from Lists. '
        'In Findings of the Association for Computational Linguistics: ACL 2022, pages 1105\u20131119.',

        'Chen, Tong, Danny Wang, Xurong Liang, Marten Risius, Gianluca Demartini, and Hongzhi Yin. 2024. '
        'Hate Speech Detection with Generalizable Target-aware Fairness. '
        'In Proceedings of the 2024 Conference on Empirical Methods in Natural Language Processing.',

        'Czarnowska, Paula, Yogarshi Vyas, and Kashif Shah. 2021. '
        'Quantifying Social Biases in NLP: A Generalization and Empirical Comparison of Extrinsic Fairness Metrics. '
        'Transactions of the Association for Computational Linguistics, 9:1249\u20131267.',

        'Dixon, Lucas, John Li, Jeffrey Sorensen, Nithum Thain, and Lucy Vasserman. 2018. '
        'Measuring and Mitigating Unintended Bias in Text Classification. '
        'In Proceedings of the 2018 AAAI/ACM Conference on AI, Ethics, and Society, pages 67\u201373.',

        'Elazar, Yanai, and Yoav Goldberg. 2018. '
        'Adversarial Removal of Demographic Attributes from Text Data. '
        'In Proceedings of the 2018 Conference on Empirical Methods in Natural Language Processing, pages 11\u201321.',

        'Garg, Sahaj, Vincent Perot, Nicole Limtiaco, Ankur Taly, Ed H. Chi, and Alex Beutel. 2019. '
        'Counterfactual Fairness in Text Classification through Robustness. '
        'In Proceedings of the 2019 AAAI/ACM Conference on AI, Ethics, and Society, pages 219\u2013226.',

        'Hartvigsen, Thomas, Saadia Gabriel, Hamid Palangi, Dipankar Ray, and Ece Kamar. 2022. '
        'ToxiGen: A Large-Scale Machine-Generated Dataset for Adversarial and Implicit Hate Speech Detection. '
        'In Proceedings of the 60th Annual Meeting of the Association for Computational Linguistics, pages 3309\u20133326.',

        'He, Pengcheng, Jianfeng Gao, and Weizhu Chen. 2021. '
        'DeBERTaV3: Improving DeBERTa using ELECTRA-Style Pre-Training with Gradient-Disentangled Embedding Sharing. '
        'arXiv preprint arXiv:2111.09543.',

        'Kannan, Harini, Alexey Kurakin, and Ian Goodfellow. 2018. '
        'Adversarial Logit Pairing. arXiv preprint arXiv:1803.06373.',

        'Khosla, Prannay, Piotr Teterwak, Chen Wang, Aaron Sarna, Yonglong Tian, Phillip Isola, '
        'Aaron Maschinot, Ce Liu, and Dilip Krishnan. 2020. '
        'Supervised Contrastive Learning. In Advances in Neural Information Processing Systems, 33:18661\u201318673.',

        'Kusner, Matt J., Joshua Loftus, Chris Russell, and Ricardo Silva. 2017. '
        'Counterfactual Fairness. In Advances in Neural Information Processing Systems, 30:4066\u20134076.',

        'Loshchilov, Ilya, and Frank Hutter. 2019. '
        'Decoupled Weight Decay Regularization. '
        'In International Conference on Learning Representations.',

        'Lu, Junyu, Bo Xu, Xiaokun Zhang, Kaiyuan Liu, Dongyu Zhang, Liang Yang, and Hongfei Lin. 2024. '
        'Take Its Essence, Discard Its Dross: A Counterfactual Causal Debiasing Framework for Hate Speech Detection. '
        'In Proceedings of the 2024 Conference on Empirical Methods in Natural Language Processing.',

        'Mathew, Binny, Punyajoy Saha, Seid Muhie Yimam, Chris Biemann, Pawan Goyal, and Animesh Mukherjee. 2021. '
        'HateXplain: A Benchmark Dataset for Explainable Hate Speech Detection. '
        'In Proceedings of the AAAI Conference on Artificial Intelligence, 35(17):14867\u201314875.',

        'Mostafazadeh Davani, Aida, Ali Omrani, Brendan Kennedy, Mohammad Atari, Xiang Ren, and Morteza Dehghani. 2021. '
        'Improving Counterfactual Generation for Fair Hate Speech Detection. '
        'In Proceedings of the 5th Workshop on Online Abuse and Harms (WOAH), pages 92\u2013101.',

        'Ramponi, Alan, and Sara Tonelli. 2022. '
        'Features or Spurious Artifacts? Data-Centric Baselines for Fair and Robust Hate Speech Detection. '
        'In Proceedings of the 2022 Conference of the North American Chapter of the Association for Computational Linguistics, pages 3027\u20133040.',

        'Sap, Maarten, Dallas Card, Saadia Gabriel, Yejin Choi, and Noah A. Smith. 2019. '
        'The Risk of Racial Bias in Hate Speech Detection. '
        'In Proceedings of the 57th Annual Meeting of the Association for Computational Linguistics, pages 1668\u20131678.',

        'Vidgen, Bertie, Tristan Thrush, Zeerak Waseem, and Douwe Kiela. 2021. '
        'Learning from the Worst: Dynamically Generated Datasets to Improve Online Hate Detection. '
        'In Proceedings of the 59th Annual Meeting of the Association for Computational Linguistics, pages 1667\u20131682.',

        'Zmigrod, Ran, Sabrina J. Mielke, Hanna Wallach, and Ryan Cotterell. 2019. '
        'Counterfactual Data Augmentation for Mitigating Gender Stereotypes in Languages with Rich Morphology. '
        'In Proceedings of the 57th Annual Meeting of the Association for Computational Linguistics, pages 1651\u20131661.',

        'ZhipuAI. 2024. GLM-4: An Open Multilingual Multimodal Chat Language Model. '
        'Technical Report. https://open.bigmodel.cn/',

        'Devlin, Jacob, Ming-Wei Chang, Kenton Lee, and Kristina Toutanova. 2019. '
        'BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. '
        'In Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics, pages 4171\u20134186.',

        'Liu, Yinhan, Myle Ott, Naman Goyal, Jingfei Du, Mandar Joshi, Danqi Chen, Omer Levy, Mike Lewis, '
        'Luke Zettlemoyer, and Veselin Stoyanov. 2019. '
        'RoBERTa: A Robustly Optimized BERT Pretraining Approach. arXiv preprint arXiv:1907.11692.',

        'Blodgett, Su Lin, Solon Barocas, Hal Daume III, and Hanna Wallach. 2020. '
        'Language (Technology) is Power: A Critical Survey of "Bias" in NLP. '
        'In Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics, pages 5454\u20135476.',

        'Wiegand, Michael, Josef Ruppenhofer, and Thomas Kleinbauer. 2019. '
        'Detection of Abusive Language: the Problem of Biased Datasets. '
        'In Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics, pages 602\u2013608.',

        'Park, Ji Ho, Jamin Shin, and Pascale Fung. 2018. '
        'Reducing Gender Bias in Abusive Language Detection. '
        'In Proceedings of the 2018 Conference on Empirical Methods in Natural Language Processing, pages 2799\u20132804.',

        'Kennedy, Brendan, Xisen Jin, Aida Mostafazadeh Davani, Morteza Dehghani, and Xiang Ren. 2020. '
        'Contextualizing Hate Speech Classifiers with Post-hoc Explanation. '
        'In Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics, pages 5005\u20135015.',

        'Badjatiya, Pinkesh, Shashank Gupta, Manish Gupta, and Vasudeva Varma. 2017. '
        'Deep Learning for Hate Speech Detection in Tweets. '
        'In Proceedings of the 26th International Conference on World Wide Web Companion, pages 759\u2013760.',

        'Hardt, Moritz, Eric Price, and Nathan Srebro. 2016. '
        'Equality of Opportunity in Supervised Learning. '
        'In Advances in Neural Information Processing Systems, 29:3315\u20133323.',
    ]

    for ref in refs:
        add_formatted_paragraph(doc, ref, size=9, space_after=3)

    return doc

# ============================================================
# MAIN
# ============================================================
def main():
    print("=" * 60)
    print("Generating EMNLP 2026 Paper Draft (Enhanced)")
    print("=" * 60)

    # Step 1: Generate document with title, abstract, intro, related work
    doc, baselines, ablation, arch_path, pipe_path = generate_paper()

    # Step 2: Method section with architecture diagram
    doc = add_method_section(doc, arch_path)

    # Step 3: Experimental setup with pipeline diagram
    doc = add_experimental_setup(doc, pipe_path)

    # Step 4: Main results with tables
    doc = add_results_section(doc, baselines, ablation)

    # Step 5: Ablation study
    doc = add_ablation_section(doc, ablation)

    # Step 6: Conclusion, limitations, ethics, references
    doc = add_conclusion_and_references(doc)

    # Save
    output_path = os.path.join(BASE_DIR, "docs", "EMNLP_2026_LLM_CLP_Final.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\n[DONE] Paper saved to: {output_path}")
    print(f"[DONE] Architecture diagram: {arch_path}")
    print(f"[DONE] Pipeline diagram: {pipe_path}")


if __name__ == "__main__":
    main()
