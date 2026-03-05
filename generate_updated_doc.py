"""
生成更新版论文初稿 Word 文档。
基于原始 '初.docx' 的内容，更新为包含完整方法论故事的顶刊级别文档。
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 全局样式设置 ==========
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_formula(text):
    """添加居中公式"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    run.italic = True
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(36 + level * 18)
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        pass
    else:
        run.text = text
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ================================================================
# 标题
# ================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    'Bias-Aware Toxic Comment Detection with DeBERTaV3:\n'
    'Multi-Task Learning, Anchor-PCGrad, and Slice-Level Fairness Optimization'
)
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()  # 空行

# ================================================================
# 1. 引言与问题定义
# ================================================================
add_heading('1  Introduction and Problem Definition', level=1)

add_para(
    '在线毒性评论检测（Toxic Comment Detection）是内容审核领域的核心任务。'
    '给定一条英文评论 comment_text，模型需输出毒性概率 p(toxic)，经阈值化后得到二分类结果：'
    '1 = 违规（toxic），0 = 不违规（non-toxic）。',
    indent=True
)

add_para(
    '然而，Jigsaw Unintended Bias 数据集揭示了一个严峻问题：训练数据中包含身份词'
    '（如 muslim, gay, black 等）的评论往往具有更高的 toxic 标注比例，'
    '导致模型容易学到"看到身份词就判违规"的捷径规则（shortcut learning），'
    '产生系统性的非预期偏差（unintended bias）。'
    '这种偏差在实际部署中会导致对特定人群的不公平审核。',
    indent=True
)

add_para(
    '本文提出一种三阶段渐进式训练框架，围绕 DeBERTa-v3-base 构建多任务学习骨干网络，'
    '并通过梯度冲突分析驱动的公平性优化模块（Anchor-PCGrad + Slice Ranking Loss + CLP），'
    '在保持高分类性能的同时显著降低身份偏差。方法设计遵循以下逻辑链条：',
    indent=True
)

add_bullet('Stage 1（多任务预训练）：通过辅助任务（身份识别 + 毒性子类型预测）提供正则化信号，帮助模型学习"身份 ≠ 毒性"的区分能力')
add_bullet('Stage 2（身份感知去偏精调）：引入 identity-aware 分组差异化重加权，专门修正高误报身份组')
add_bullet('Stage 3（公平性优化精调）：冻结辅助任务头，引入 Slice Ranking Loss 直接优化子群 AUC、'
           'Anchor-PCGrad 解决梯度冲突、CLP 约束反事实一致性')

# ================================================================
# 2. 数据集与预处理
# ================================================================
add_heading('2  Dataset and Preprocessing', level=1)

add_heading('2.1  数据来源', level=2)

add_para(
    '实验数据来自 Jigsaw Unintended Bias in Toxicity Classification（Kaggle, 2019）。'
    '原始训练集约 180 万条英文评论，包含以下标注信息：',
    indent=True
)

add_bullet('target（0~1 连续值）：毒性程度的众包标注均值，≥0.5 视为 toxic')
add_bullet('6 种毒性子类型（0~1）：severe_toxicity, obscene, threat, insult, identity_attack, sexual_explicit')
add_bullet('9 种身份标签（0~1）：male, female, black, white, muslim, jewish, christian, homosexual_gay_or_lesbian, psychiatric_or_mental_illness')

add_heading('2.2  数据预处理流程', level=2)

add_para(
    '为保证实验的可复现性和公平性，数据预处理采用以下严格流程：',
    indent=True
)

add_para('(1) 文本清洗', bold=True)
add_bullet('保留原始文本，不进行大小写转换或特殊符号移除（DeBERTa tokenizer 自带 subword 分词能力）')
add_bullet('截断长度 max_len = 256 tokens（覆盖 >95% 样本）')

add_para('(2) 标签处理', bold=True)
add_bullet('主任务 y_tox：直接使用 target 原始连续值作为 Soft Label（clip 到 [0,1]），保留标注不确定性信息。'
           '二分类判定阈值 ≥ 0.5')
add_bullet('子类型 y_sub：6 维向量，NaN 填 0，保留 [0,1] 软标签')
add_bullet('身份 y_id：9 维向量，NaN 填 0，保留 [0,1] 软标签')
add_bullet('has_identity：若 max(y_id) ≥ 0.5 则为 1，否则为 0。用于决定是否施加身份重加权')

add_para('(3) 数据采样与划分', bold=True)
add_bullet('从原始 180 万条中采样 30 万条（seed=42），确保 toxic/non-toxic 比例均衡（1:1）')
add_bullet('划分比例：train : val : test = 8 : 1 : 1')
add_bullet('采样使用固定 data_seed=42，与模型训练 seed 解耦，保证不同 seed 实验使用完全相同的数据子集')
add_bullet('预处理后数据以 Parquet 格式持久化存储（train_processed.parquet, val_processed.parquet, test_processed.parquet）')

add_para('(4) 数据统计', bold=True)

# 添加数据统计表
table = doc.add_table(rows=5, cols=4, style='Table Grid')
headers = ['Split', 'Total', 'Toxic (≥0.5)', 'Non-toxic (<0.5)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for para in table.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True

data_rows = [
    ['Train', '~223,000', '~111,600', '~111,400'],
    ['Val', '~28,000', '~14,000', '~14,000'],
    ['Test', '~28,000', '~14,000', '~14,000'],
    ['Total', '~300,000', '~150,000', '~150,000'],
]
for i, row in enumerate(data_rows):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ================================================================
# 3. 模型架构
# ================================================================
add_heading('3  Model Architecture', level=1)

add_heading('3.1  Backbone: DeBERTa-v3-base Encoder', level=2)

add_para(
    '选用 DeBERTa-v3-base（He et al., 2021）作为文本编码器。相较于 BERT 和 RoBERTa，'
    'DeBERTa-v3 具有两项关键改进：',
    indent=True
)

add_bullet('Disentangled Attention：将 token 的内容向量和位置向量解耦，'
           '分别计算 content-to-content、content-to-position、position-to-content 三种注意力分量，'
           '提升对相对位置信息的建模能力')
add_bullet('Replaced Token Detection (RTD)：采用 ELECTRA 风格的预训练目标替代 MLM，'
           '配合 gradient-disentangled embedding sharing (GDES)，在相同参数量下获得更强的语言理解能力')

add_para(
    '输入经 DeBERTa tokenizer 编码为 input_ids 和 attention_mask，'
    '编码器输出最后一层 token 表示 H = [h₁, h₂, ..., h_L]，形状 [B, L, d]（d=768）。',
    indent=True
)

add_heading('3.2  Pooling: CLS + Attention Pooling (Tanh Bottleneck)', level=2)

add_para(
    '毒性触发通常集中在评论中的特定词汇或短语。仅使用 [CLS] 向量可能遗漏局部触发信号。'
    '因此，我们采用 CLS 向量与 Attention Pooling 的拼接策略：',
    indent=True
)

add_para('CLS 向量：h_cls = H[:,0,:]（全局语义表示）', bold=False)
add_para('Attention Pooling（Tanh Bottleneck 结构）：', bold=True)
add_formula('score = Linear(d → d/4) → Tanh → Linear(d/4 → 1)')
add_formula('α = softmax(masked_fill(score, padding → -∞))')
add_formula('h_att = Σ αᵢ · hᵢ')

add_para(
    '拼接特征：h = concat(h_cls, h_att) ∈ ℝ^{2d}。'
    'CLS 捕获全局语义，Attention Pooling 聚焦局部攻击信号，互补提升鲁棒性。'
    'Tanh Bottleneck 结构（d → d/4 → 1）相比简单的 Linear(d → 1) 提供更强的非线性拟合能力。',
    indent=True
)

add_heading('3.3  Three-Task Multi-Head with Physical Isolation', level=2)

add_para(
    '为支持多任务学习同时避免梯度冲突，三个任务头采用完全物理隔离的投影层设计。'
    '每个任务拥有独立的 projection → activation → dropout → head 路径，'
    '不共享任何参数：',
    indent=True
)

add_bullet('主任务路径：z_tox = Dropout₀.₂(GELU(Linear(h, 512)))  →  logit_tox = Linear(z_tox, 1)')
add_bullet('子类型路径：z_sub = Dropout₀.₁(GELU(Linear(h, 512)))  →  logit_sub = Linear(z_sub, 6)')
add_bullet('身份路径：  z_id  = Dropout₀.₁(GELU(Linear(h, 512)))  →  logit_id  = Linear(z_id, 9)')

add_para(
    '推理时仅使用主任务头：p_tox = σ(logit_tox)。辅助任务头在 Stage 3 中被冻结并重新定位为'
    '"切片定义器"（Slice Definer），用于确定身份子群的成员关系，而非直接参与损失优化。',
    indent=True
)

add_heading('3.4  Uncertainty Weighting (Kendall et al., 2018)', level=2)

add_para(
    '多任务损失加权采用 homoscedastic uncertainty weighting 自动学习：',
    indent=True
)

add_formula('L_total = L_tox / (2·exp(σ_tox)) + σ_tox/2 + L_sub / (2·exp(σ_sub)) + σ_sub/2 + L_id / (2·exp(σ_id)) + σ_id/2')

add_para(
    '其中 σ_tox, σ_sub, σ_id 为可学习参数（初始化为 0）。不确定性高的任务自动获得较低权重，'
    '无需手动调节 α、β 等超参数。在 Stage 3 中，这三个参数同样被冻结。',
    indent=True
)

# ================================================================
# 4. 方法论：从梯度冲突到公平性优化
# ================================================================
add_heading('4  Methodology: From Gradient Conflict to Fairness Optimization', level=1)

add_para(
    '本节阐述方法设计的核心逻辑链条。我们首先分析多任务学习框架中存在的梯度冲突问题，'
    '然后逐一引入三个针对性的解决模块。',
    indent=True
)

add_heading('4.1  Motivation: 多任务梯度冲突分析', level=2)

add_para(
    '在 Stage 2 的 identity-aware 重加权训练中，我们观察到一个关键现象：'
    '尽管整体 ROC-AUC 可以通过调参提升（如 ParamFix 实验的 AUC=0.9708），'
    '但偏见相关指标（Subgroup AUC, BPSN AUC, BNSP AUC）并未同步改善，'
    '甚至出现退化。在 Jigsaw 官方 Final Metric 下，AUC 更高的 ParamFix 模型'
    '反而劣于原始 MTL 模型（0.9332 vs 0.9391）。',
    indent=True
)

add_para(
    '分析原因：主任务损失梯度（优化整体分类能力）与公平性损失梯度（缩小子群间差异）'
    '在 backbone 层面存在方向冲突。直接加权求和会导致两种梯度相互抵消，'
    '使优化器在"提升性能"和"改善公平"之间震荡，无法同时达到最优。',
    indent=True
)

add_para(
    '这一发现促使我们设计 Stage 3 的公平性优化框架。核心思路：'
    '不再让公平性信号通过辅助任务头的梯度回传影响 backbone，'
    '而是将辅助任务头冻结，仅利用其输出定义"身份切片"（identity slice），'
    '在此基础上施加专门的公平性损失，并通过 PCGrad 机制解决梯度冲突。',
    indent=True
)

add_heading('4.2  Slice Ranking Loss: 直接优化子群 AUC', level=2)

add_para(
    'Jigsaw 官方评估指标 Final Metric 的核心是 Subgroup AUC、BPSN AUC 和 BNSP AUC。'
    '传统的 BCE 损失是逐样本的点式损失（pointwise loss），无法直接优化排序质量（AUC）。'
    '我们引入 Slice Ranking Loss，将三种子群 AUC 转化为可微的成对排序损失（pairwise ranking loss）：',
    indent=True
)

add_para('对于每个身份组 g ∈ {1, ..., G}，构造三组正/负样本对：', bold=True)

add_bullet('Subgroup 对: 组内正样本 vs 组内负样本 → 优化 Subgroup AUC')
add_bullet('BPSN 对: 背景正样本 vs 组内负样本 → 优化 BPSN AUC（减少子群无毒样本被误判）')
add_bullet('BNSP 对: 组内正样本 vs 背景负样本 → 优化 BNSP AUC（减少子群有毒样本被漏判）')

add_para('对每组正/负对，使用 pairwise logistic loss 作为 AUC 的可微代理：', bold=True)
add_formula('ℓ(p, n) = log(1 + exp(-(s_p - s_n)))')

add_para(
    '其中 s_p 和 s_n 分别为正/负样本的 logit。当正样本得分高于负样本时，损失趋近于 0；'
    '反之损失增大，驱动模型正确排序。',
    indent=True
)

add_para(
    '各组的损失通过 Power Mean（p=4）聚合，赋予表现最差的子群更高权重：',
    indent=True
)

add_formula('L_slice = (1/G · Σ L_g^p)^{1/p},  p = 4')

add_para(
    '与 Final Metric 使用 p = -5 强调最差组不同，训练时使用 p = 4（正值）'
    '以避免梯度爆炸，同时仍能对高损失组施加更大优化压力。',
    indent=True
)

add_para(
    '效率优化：当正负样本对数量超过 max_pairs=256 时，采用随机采样而非全组合，'
    '将计算复杂度从 O(n²) 降至 O(max_pairs)，保证大 batch 下的训练效率。',
    indent=True
)

add_heading('4.3  Anchor-PCGrad: 非对称梯度冲突消解', level=2)

add_para(
    'PCGrad（Yu et al., 2020）是一种多任务梯度冲突消解方法，'
    '当两个任务梯度方向冲突（内积 < 0）时，将一个任务的梯度投影到另一个的法平面上。'
    '但标准 PCGrad 是对称的——两个梯度都可能被修改，在我们的场景中不合理：'
    '主任务（毒性检测性能）是刚性约束，不应被公平性信号削弱。',
    indent=True
)

add_para(
    '因此我们提出 Anchor-PCGrad（锚定式梯度冲突消解），核心原则为"主任务梯度绝不修改"：',
    indent=True
)

add_para('算法流程：', bold=True)
add_bullet('Step 1: 对主任务损失 L_main 单独反向传播，获取梯度 g_main')
add_bullet('Step 2: 对公平性损失 L_bias（= L_slice + λ_clp · L_clp）单独反向传播，获取梯度 g_bias')
add_bullet('Step 3: 检测 backbone 层上的梯度冲突：dot = ⟨g_bias, g_main⟩')
add_bullet('Step 4: 若 dot < 0（冲突），将 g_bias 投影到 g_main 的法平面：')

add_formula("g_bias' = g_bias - (⟨g_bias, g_main⟩ / ||g_main||²) · g_main")

add_bullet('Step 5: 合并最终梯度：g_final = g_main + λ_bias · g_bias\'')

add_para(
    '关键设计：冲突检测和投影仅在 backbone 参数上进行（排除 proj_tox、tox_head 等任务特定参数），'
    '因为梯度冲突主要发生在共享表示层。任务头层面的梯度可以自由叠加。',
    indent=True
)

add_para(
    'DDP 兼容性：在 Anchor-PCGrad 过程中使用 model.no_sync() 上下文管理器'
    '禁止自动梯度同步，待合并最终梯度后手动执行 all_reduce 均值归约。',
    indent=True
)

add_heading('4.4  CLP: Counterfactual Logit Pairing', level=2)

add_para(
    '反事实公平性（Counterfactual Fairness）的核心直觉：如果一条评论仅因提到了某个身份词'
    '而被判为 toxic，这种判定就是不公平的。理想状态下，具有相同毒性标签的样本'
    '（无论是否提到身份词）应获得相近的预测 logit。',
    indent=True
)

add_para(
    '受 Garg et al. (2019) 的 Counterfactual Logit Pairing 启发，'
    '我们实现了一种 batch 级别的 CLP 近似。核心思想：'
    '在每个 mini-batch 内，对所有提到身份词的样本，惩罚同标签样本间的 logit 差异：',
    indent=True
)

add_formula('L_clp = Σᵢⱼ wᵢⱼ · (sᵢ - sⱼ)² / Σᵢⱼ wᵢⱼ')

add_para('其中权重矩阵为：', bold=True)

add_bullet('wᵢⱼ = 1（同标签对，只考虑上三角避免重复）')
add_bullet('wᵢⱼ × 3.0（若两个样本均为 non-toxic，额外提升权重——这对应 BPSN 场景，'
           '即"提到身份的无毒样本不应被误判"）')
add_bullet('wᵢⱼ = 0（不同标签对，不施加约束）')

add_para(
    '此设计无需真正的文本替换（text-level CDA），而是利用 batch 内的自然变异'
    '构造"伪反事实"对，计算代价极低且与 DDP 天然兼容。',
    indent=True
)

# ================================================================
# 5. 三阶段训练流程
# ================================================================
add_heading('5  Three-Stage Training Pipeline', level=1)

add_heading('5.1  Stage 1: Multi-Task Pre-training', level=2)

add_para('目标：让模型同时学会毒性语义 + 子类型预测 + 身份属性识别，三任务互相正则化。', bold=True)

add_bullet('损失函数：Focal Loss (α=1.0, γ=2.0) 用于主任务，BCEWithLogitsLoss 用于辅助任务')
add_bullet('多任务加权：Uncertainty Weighting 自动学习')
add_bullet('训练配置：8 卡 DDP + AMP (FP16)，单卡 batch=16，梯度累积 2 步（等效 batch=256）')
add_bullet('优化器：AdamW (lr=1e-5, weight_decay=0.01)')
add_bullet('学习率调度：线性预热（10% 总步数）+ 线性衰减至 0')
add_bullet('Epochs: 6')

add_heading('5.2  Stage 2: Identity-Aware Debiasing Fine-tuning', level=2)

add_para('目标：在 S1 已学好的基础上，重点修正"提到身份词但非毒性"的误报。', bold=True)

add_bullet('加载 S1 最优 checkpoint')
add_bullet('主任务损失：identity-aware weighted BCE，权重规则：')
add_bullet('  has_identity=1 且 non-toxic: w = max(该样本涉及的身份组权重)，'
           '权重表: female=1.0, male=1.2, christian=1.3, white=1.8, muslim=1.9, '
           'black=2.3, homosexual=2.6, jewish=2.9, psychiatric=3.4')
add_bullet('  has_identity=1 且 toxic: w = 1.5（固定）')
add_bullet('  其他样本: w = 1.0')
add_bullet('Layer-wise Learning Rate Decay (decay=0.95)：底层保持稳定，顶层适度调整')
add_bullet('训练配置：8 卡 DDP + AMP，单卡 batch=96，梯度累积 4 步（等效 batch=3072）')
add_bullet('优化器：AdamW (lr=3e-6), Epochs: 4, EMA (decay=0.999)')

add_heading('5.3  Stage 3: Fair S2 — Fairness-Aware Optimization', level=2)

add_para(
    '目标：在保持 S2 分类性能的基础上，通过直接优化子群 AUC 相关的排序损失 + 梯度冲突消解 + '
    '反事实一致性约束，最大化 Jigsaw Official Final Metric。',
    bold=True
)

add_para('核心改动（相对于 S2）：', bold=True)

add_bullet('冻结辅助任务头：proj_sub, proj_id, subtype_head, identity_head, log_var 参数全部 requires_grad=False')
add_bullet('新增 Slice Ranking Loss：直接优化 Subgroup/BPSN/BNSP AUC 排序质量')
add_bullet('新增 Anchor-PCGrad：解决 L_main 与 L_bias 的 backbone 梯度冲突')
add_bullet('新增 CLP Loss：约束同标签身份样本的 logit 一致性')

add_para('三阶段训练调度（Phase Schedule）：', bold=True)

# 添加调度表
table2 = doc.add_table(rows=4, cols=4, style='Table Grid')
h2 = ['Phase', 'Progress Range', 'λ_bias', 'λ_clp']
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
    for para in table2.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True

phase_data = [
    ['Warmup', '0 ~ 20%', '0 (仅主任务)', '0'],
    ['Debias', '20% ~ 90%', '0 → λ_bias (线性增长)', '0 → λ_clp (线性增长)'],
    ['Stabilize', '90% ~ 100%', 'λ_bias × 0.5', 'λ_clp × 0.5'],
]
for i, row in enumerate(phase_data):
    for j, val in enumerate(row):
        table2.rows[i+1].cells[j].text = val

doc.add_paragraph()

add_para(
    '设计理由：Warmup 阶段先让模型适应新学习率，避免公平性损失在初始不稳定阶段干扰训练；'
    'Debias 阶段线性增加公平性损失权重，让梯度冲突消解有缓冲；'
    'Stabilize 阶段降低公平性损失权重，让模型在接近收敛时稳定。',
    indent=True
)

add_para('Fair S2 训练超参数：', bold=True)
add_bullet('lr = 5e-6, epochs = 6, batch_size = 48/卡（6卡 DDP）')
add_bullet('λ_bias = 1.0, λ_clp = 0.1, power_p = 4')
add_bullet('EMA decay = 0.999, Layer-wise LR decay = 0.95')
add_bullet('Cosine Annealing 学习率调度')
add_bullet('模型选择标准：以验证集 Final Metric 为准（而非 F1 或 AUC）')

# ================================================================
# 6. 评估指标体系
# ================================================================
add_heading('6  Evaluation Metrics', level=1)

add_para(
    '本任务采用 Jigsaw Unintended Bias Competition（Borkan et al., 2019）定义的官方评估框架，'
    '以 Final Metric 作为模型选择的首要指标。评估体系分三个层次：',
    indent=True
)

add_heading('6.1  主任务分类指标', level=2)
add_bullet('F1 Score（正类 = toxic）：使用阈值扫描（0.01 ~ 0.99）选取最优阈值')
add_bullet('Accuracy：整体正确率')
add_bullet('ROC-AUC：阈值无关的排序能力度量')
add_bullet('PR-AUC：在类别不均衡时比 ROC-AUC 更敏感')

add_heading('6.2  偏见/群体鲁棒指标 (Nuanced Metrics)', level=2)

add_para(
    '对 9 个身份子群，分别计算三种 AUC（Borkan et al., 2019）：',
    indent=True
)

add_bullet('Subgroup AUC: 仅在提到身份 g 的子集上计算 AUC')
add_bullet('BPSN AUC (Background Positive, Subgroup Negative): 背景正样本 ∪ 子群负样本的 AUC')
add_bullet('BNSP AUC (Background Negative, Subgroup Positive): 背景负样本 ∪ 子群正样本的 AUC')

add_para('聚合指标：', bold=True)
add_bullet('Mean Bias AUC: 所有子群 × 3 种 AUC 的算术均值')
add_bullet('Worst-group Bias AUC: 所有子群 × 3 种 AUC 的最小值')

add_heading('6.3  Jigsaw Official Final Metric（首要指标）', level=2)

add_formula('Final = 0.25 × Overall_AUC + 0.75 × BiasScore')

add_para('其中 BiasScore 使用 Generalized Mean（p = -5）聚合，强调最差子群表现：', indent=True)

add_formula('BiasScore = 1/3 × (M_{-5}(Subgroup_AUCs) + M_{-5}(BPSN_AUCs) + M_{-5}(BNSP_AUCs))')

add_formula('M_p(x₁, ..., x_n) = (1/n · Σ xᵢ^p)^{1/p},  p = -5')

add_para(
    'p = -5 使得当任何一个子群的 AUC 极低时，BiasScore 会被严重拉低，'
    '从而惩罚"整体好但个别子群差"的模型。Final Metric 中 BiasScore 权重占 75%，'
    '体现了对公平性的高要求。',
    indent=True
)

# ================================================================
# 7. 实验设计
# ================================================================
add_heading('7  Experimental Design', level=1)

add_heading('7.1  对比实验模型 (Baselines)', level=2)

add_bullet('TF-IDF + Logistic Regression: 传统文本分类强基线')
add_bullet('BERT-base fine-tune: 预训练 Transformer 基线')
add_bullet('RoBERTa-base fine-tune: 预训练 Transformer 基线')
add_bullet('BERT + CNN-BiLSTM: 非预训练深度学习基线')
add_bullet('DeBERTa-v3-base (Vanilla): 单任务 DeBERTa 基线（无多任务、无重加权）')
add_bullet('DeBERTa-v3 MTL (S1+S2): 多任务 + identity-aware 重加权（无 Stage 3）')
add_bullet('DeBERTa-v3 MTL + Fair S2 (Full): 完整三阶段框架（本文方法）')

add_heading('7.2  消融实验 (Ablation Study)', level=2)

add_para(
    '为验证各组件的贡献，设计以下消融实验（每组仅移除一个因素）：',
    indent=True
)

add_para('A. 基础架构消融（在 S1+S2 基础上）：', bold=True)
add_bullet('NoPooling: 移除 Attention Pooling，仅使用 CLS 向量')
add_bullet('NoFocal: S1 使用标准 BCE 替代 Focal Loss')
add_bullet('NoReweight: S2 不使用身份重加权（所有样本 w=1）')
add_bullet('OnlyToxicity: 移除所有辅助任务（纯单任务训练）')

add_para('B. Fair S2 组件消融：', bold=True)
add_bullet('NoSlice: 移除 Slice Ranking Loss（仅保留 PCGrad + CLP）')
add_bullet('NoPCGrad: 移除 Anchor-PCGrad（直接梯度叠加）')
add_bullet('NoCLP: 移除 CLP Loss（仅保留 Slice + PCGrad）')

add_heading('7.3  超参敏感性分析', level=2)

add_bullet('身份重加权强度: w_identity_toxic = {1.0, 1.5, 2.0}')
add_bullet('公平性损失权重: λ_bias = {0.5, 1.0, 2.0}')
add_bullet('CLP 损失权重: λ_clp = {0.05, 0.1, 0.2}')

add_heading('7.4  统计显著性', level=2)

add_para(
    '所有主要实验使用 3 个随机种子（42, 123, 2024）重复，报告均值 ± 标准差。'
    '模型训练 seed 与数据采样 seed（固定为 42）解耦，确保不同 seed 实验使用完全相同的数据。',
    indent=True
)

# ================================================================
# 8. 方法总结
# ================================================================
add_heading('8  Method Summary', level=1)

add_para(
    '本文提出的完整方法可概括为以下技术链条：',
    indent=True
)

add_bullet('① DeBERTa-v3-base Encoder + CLS/Attention Pooling (Tanh Bottleneck)')
add_bullet('② 物理隔离的三任务投影层 proj_tox / proj_sub / proj_id + 三个输出头')
add_bullet('③ Focal Loss + Uncertainty Weighting 自动多任务权重（Stage 1）')
add_bullet('④ Soft Label 训练（保留连续标注值 [0,1] 的不确定性信息）')
add_bullet('⑤ 两阶段基础训练：S1 多任务预训练 → S2 identity-aware 分组差异化重加权去偏精调')
add_bullet('⑥ 梯度冲突分析 → 发现主任务与公平性目标的梯度对抗')
add_bullet('⑦ Fair S2 公平性优化：Slice Ranking Loss + Anchor-PCGrad + CLP（Stage 3）')
add_bullet('⑧ 三相训练调度：Warmup → Debias (λ ramp-up) → Stabilize')
add_bullet('⑨ 以 Jigsaw Official Final Metric 作为模型选择首要标准')

# ================================================================
# 参考文献
# ================================================================
add_heading('References', level=1)

refs = [
    'Borkan, D., Dixon, L., Sorensen, J., Thain, N., & Vasserman, L. (2019). Nuanced Metrics for Measuring Unintended Bias with Real Data for Text Classification. In Companion Proceedings of The 2019 World Wide Web Conference (WWW).',
    'Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. In NAACL-HLT.',
    'Garg, S., Perot, V., Limtiaco, N., Taly, A., Chi, E. H., & Beutel, A. (2019). Counterfactual Fairness in Text Classification through Robustness. In AIES.',
    'He, P., Liu, X., Gao, J., & Chen, W. (2020). DeBERTa: Decoding-enhanced BERT with Disentangled Attention. In ICLR 2021.',
    'He, P., Gao, J., & Chen, W. (2021). DeBERTaV3: Improving DeBERTa using ELECTRA-Style Pre-Training with Gradient-Disentangled Embedding Sharing. arXiv:2111.09543.',
    'Kendall, A., Gal, Y., & Cipolla, R. (2018). Multi-Task Learning Using Uncertainty to Weigh Losses for Scene Geometry and Semantics. In CVPR.',
    'Kim, Y. (2014). Convolutional Neural Networks for Sentence Classification. In EMNLP.',
    'Lin, T.-Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal Loss for Dense Object Detection. In ICCV.',
    'Liu, Y., Ott, M., Goyal, N., et al. (2019). RoBERTa: A Robustly Optimized BERT Pretraining Approach. arXiv:1907.11692.',
    'Yu, T., Kumar, S., Gupta, A., Levine, S., Hausman, K., & Finn, C. (2020). Gradient Surgery for Multi-Task Learning. In NeurIPS.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(-22)
    p.paragraph_format.left_indent = Pt(22)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

# ========== 保存 ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '初_v2_FairS2.docx')
doc.save(output_path)
print(f"文档已保存: {output_path}")
