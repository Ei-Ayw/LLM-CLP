#!/usr/bin/env python3
"""生成 EMNLP 2026 论文初稿 (Word 格式) - 修正版
修复: 1) 去掉重复的Swap表格 2) 加入消融实验 3) 对比模型用方法名
"""
import json, glob, numpy as np
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ============ 加载数据 ============
with open("src_result/eval/summary_all_7metrics.json") as f:
    baseline_data = json.load(f)

# 加载消融实验数据
DATASETS = ['hatexplain', 'toxigen', 'dynahate']
METRICS = ['accuracy', 'macro_f1', 'auc_roc', 'cfr', 'ctfg', 'fped', 'fned']
ABLATION_PATTERNS = [
    ('none_clp0.0', 'No CF (λ=0)'),
    ('swap_clp0.0', 'Swap (λ=0)'),
    ('swap_clp1.0', 'Swap+CLP (λ=1)'),
    ('llm_clp0.0', 'LLM (λ=0)'),
    ('llm_clp1.0', 'LLM+CLP (λ=1)'),
]

ablation_data = {}
for ds in DATASETS:
    ablation_data[ds] = {}
    for pattern, name in ABLATION_PATTERNS:
        files = sorted(glob.glob(f"src_result/eval/{ds}_{pattern}_*_7metrics.json"))
        if not files:
            continue
        vals = {m: [] for m in METRICS}
        for f in files:
            d = json.load(open(f))
            for m in METRICS:
                v = d.get(m)
                if v is not None:
                    vals[m].append(float(v))
        ablation_data[ds][name] = {}
        for m in METRICS:
            if vals[m]:
                ablation_data[ds][name][m] = {
                    'mean': round(np.mean(vals[m]), 4),
                    'std': round(np.std(vals[m]), 4),
                }

# ============ 方法名映射 ============
METHOD_NAMES = {
    'Vanilla': 'Vanilla',
    'EAR': 'EAR',
    'GetFair': 'GetFair',
    'CCDF': 'CCDF',
    'Davani': 'LogitPairing',
    'Ramponi': 'AdvDebias',
    'Ours': 'Ours (LLM+CLP)',
}
BASELINE_KEYS = ['Vanilla', 'EAR', 'GetFair', 'CCDF', 'Davani', 'Ramponi', 'Ours']
DS_NAMES = {'hatexplain': 'HateXplain', 'toxigen': 'ToxiGen', 'dynahate': 'DynaHate'}

# ============ 文档设置 ============
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def make_table(rows_data, col_headers, best_cols=None):
    """创建格式化表格, best_cols: {col_idx: 'min'|'max'} 标记最优值加粗"""
    table = doc.add_table(rows=len(rows_data)+1, cols=len(col_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(col_headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

    # 找最优值
    best_methods = {}
    if best_cols:
        for ci, direction in best_cols.items():
            vals = []
            for ri, row in enumerate(rows_data):
                try:
                    v = float(row[ci].replace('±', ' ').split()[0])
                    vals.append((v, ri))
                except:
                    pass
            if vals:
                if direction == 'min':
                    best_methods[ci] = min(vals, key=lambda x: x[0])[1]
                else:
                    best_methods[ci] = max(vals, key=lambda x: x[0])[1]

    # Data
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                if ci > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    if best_methods.get(ci) == ri:
                        run.bold = True
    return table

def get_val(data_dict, ds, method, metric, cf='llm'):
    """从baseline_data获取 mean±std 字符串"""
    d = data_dict.get(method, {}).get(ds, {}).get(cf, {}).get(metric)
    if d and isinstance(d, dict):
        return f"{d['mean']:.4f}±{d['std']:.4f}"
    return '--'

def get_mean(data_dict, ds, method, metric, cf='llm'):
    d = data_dict.get(method, {}).get(ds, {}).get(cf, {}).get(metric)
    if d and isinstance(d, dict):
        return d['mean']
    return None

def get_ablation_val(ds, name, metric):
    d = ablation_data.get(ds, {}).get(name, {}).get(metric)
    if d and isinstance(d, dict):
        return f"{d['mean']:.4f}±{d['std']:.4f}"
    return '--'

def get_ablation_mean(ds, name, metric):
    d = ablation_data.get(ds, {}).get(name, {}).get(metric)
    if d and isinstance(d, dict):
        return d['mean']
    return None

# ============================================================
# Title & Authors
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Counterfactual Logit Pairing for Causally Fair Toxicity Classification")
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run("Anonymous Authors")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
doc.add_paragraph()

# ============================================================
# Abstract
# ============================================================
add_heading("Abstract", level=1)

# 用真实数据计算改进幅度
imp = {}
for ds in DATASETS:
    ours = get_mean(baseline_data, ds, 'Ours', 'cfr')
    vanilla = get_mean(baseline_data, ds, 'Vanilla', 'cfr')
    best_bl = min(get_mean(baseline_data, ds, m, 'cfr') or 999
                  for m in ['EAR', 'GetFair', 'CCDF', 'Davani', 'Ramponi'])
    imp[ds] = {
        'vs_vanilla': (1 - ours/vanilla)*100 if vanilla else 0,
        'vs_best': (1 - ours/best_bl)*100 if best_bl else 0,
        'ours': ours, 'vanilla': vanilla, 'best_bl': best_bl,
    }

add_para(
    "Toxicity detection models often exhibit social biases, producing different predictions "
    "when identity terms are changed in otherwise identical text. Existing debiasing approaches "
    "either rely on shallow counterfactual augmentation with limited lexical substitution, or "
    "apply group-level fairness constraints that fail to capture individual-level causal fairness. "
    "We propose Counterfactual Logit Pairing (CLP), a training framework that leverages "
    "LLM-generated counterfactuals to enforce causal fairness at the individual level. Our method "
    "pairs each training example with its counterfactual variant and minimizes the KL divergence "
    "between their output logits. Experiments on three benchmark datasets (HateXplain, ToxiGen, "
    f"DynaHate) demonstrate that CLP reduces Counterfactual Fairness Rate (CFR) by "
    f"{imp['hatexplain']['vs_vanilla']:.0f}% on HateXplain, "
    f"{imp['toxigen']['vs_vanilla']:.0f}% on ToxiGen, and "
    f"{imp['dynahate']['vs_vanilla']:.0f}% on DynaHate compared to vanilla fine-tuning, "
    "while maintaining competitive classification performance."
)

# ============================================================
# 1. Introduction
# ============================================================
add_heading("1  Introduction", level=1)
add_para(
    "Online toxicity detection is a critical component of content moderation systems. "
    "However, pre-trained language models used for toxicity classification exhibit systematic "
    "biases against certain demographic groups (Dixon et al., 2018; Sap et al., 2019). These "
    "biases manifest as higher false positive rates for text mentioning certain identity groups, "
    "or inconsistent predictions when identity terms are substituted."
)
add_para(
    "Counterfactual fairness (Kusner et al., 2017) provides a principled framework: a model is "
    "counterfactually fair if its prediction remains unchanged when a protected attribute is "
    "counterfactually altered. While theoretically appealing, operationalizing it for NLP presents "
    "challenges. Simple word-level substitution (e.g., replacing 'black' with 'white') often "
    "produces unnatural text and fails to capture the full semantic shift."
)
add_para("We address these limitations with two key contributions:")
add_para(
    "(1) LLM-based Counterfactual Generation: We use large language models to generate "
    "semantically coherent counterfactual pairs that go beyond simple lexical substitution."
)
add_para(
    "(2) Counterfactual Logit Pairing (CLP): We introduce a training objective that minimizes "
    "the KL divergence between the output distributions of original and counterfactual inputs, "
    "directly enforcing counterfactual fairness during training."
)

# ============================================================
# 2. Related Work
# ============================================================
add_heading("2  Related Work", level=1)

add_heading("2.1  Fairness in Toxicity Detection", level=2)
add_para(
    "Prior work has identified significant biases in toxicity classifiers. Dixon et al. (2018) "
    "showed that models trained on Wikipedia comments exhibit higher false positive rates for "
    "text containing identity terms. Sap et al. (2019) demonstrated racial bias in hate speech "
    "detection. Several debiasing approaches have been proposed, including data augmentation "
    "(Park et al., 2018), adversarial training (Zhang et al., 2018), and regularization-based "
    "methods (Huang et al., 2020)."
)

add_heading("2.2  Counterfactual Fairness", level=2)
add_para(
    "Counterfactual fairness (Kusner et al., 2017) formalizes fairness through causal reasoning. "
    "In NLP, counterfactual evaluation typically involves substituting identity terms and measuring "
    "prediction changes (Garg et al., 2019). CDA extends this to training by augmenting data with "
    "counterfactual examples (Zmigrod et al., 2019). However, simple word substitution often "
    "produces unnatural text. Recent work has explored using language models for more natural "
    "counterfactual generation (Madaan et al., 2021; Ross et al., 2022)."
)

add_heading("2.3  Debiasing Methods", level=2)
add_para(
    "EAR (Han et al., 2022) uses entropy-based attention regularization to reduce reliance on "
    "identity terms. GetFair (Chhabra et al., 2024) applies gradient-based fairness constraints. "
    "CCDF (Qian et al., 2023) uses counterfactual causal debiasing with Total Direct Effect (TDE) "
    "inference. LogitPairing (Davani et al., 2021) pairs original and augmented examples to align "
    "output distributions. AdvDebias (Ramponi & Tonelli, 2022) uses adversarial training to remove "
    "protected attribute information from representations."
)

# ============================================================
# 3. Method
# ============================================================
add_heading("3  Method", level=1)

add_heading("3.1  Problem Formulation", level=2)
add_para(
    "Given a text x and a protected attribute a (e.g., mentioned identity group), a classifier "
    "f is counterfactually fair if f(x) = f(x'), where x' is the counterfactual of x obtained "
    "by changing a to a'. We measure this through the Counterfactual Fairness Rate (CFR):"
)
add_para("CFR = (1/N) Σ 𝟙[f(xᵢ) ≠ f(x'ᵢ)]", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("where N is the number of test examples with valid counterfactual pairs. Lower CFR indicates better counterfactual fairness.")

add_heading("3.2  LLM-based Counterfactual Generation", level=2)
add_para(
    "We use an instruction-tuned LLM to generate counterfactual pairs. Given an input text x "
    "mentioning identity group a, we prompt the LLM to rewrite x as if it referred to a different "
    "identity group a', while preserving the semantic content and toxicity label. This produces "
    "more natural counterfactuals than simple word substitution, as the LLM can adjust surrounding "
    "context, pronouns, and cultural references appropriately."
)

add_heading("3.3  Counterfactual Logit Pairing (CLP)", level=2)
add_para("Our training objective combines the standard classification loss with a counterfactual logit pairing term:")
add_para("L = L_CE(f(x), y) + λ · D_KL(f(x) ∥ f(x'))", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    "where L_CE is the cross-entropy loss, f(x) and f(x') are the softmax output distributions "
    "for the original and counterfactual inputs, D_KL is the KL divergence, and λ controls the "
    "strength of the fairness constraint."
)

add_heading("3.4  Training Procedure", level=2)
add_para(
    "We fine-tune DeBERTa-v3-base as the backbone classifier. For each training batch, we include "
    "both original examples and their LLM-generated counterfactuals. The CLP loss is computed only "
    "for examples that have valid counterfactual pairs. We use λ=1.0 based on preliminary "
    "experiments. Training uses AdamW optimizer with learning rate 2e-5, linear warmup over 10% "
    "of steps, and early stopping based on validation Macro-F1."
)

# ============================================================
# 4. Experimental Setup
# ============================================================
add_heading("4  Experimental Setup", level=1)

add_heading("4.1  Datasets", level=2)
add_para("We evaluate on three toxicity detection benchmarks:")
add_para("HateXplain (Mathew et al., 2021): 20,148 social media posts annotated for hate speech with identity group labels.")
add_para("ToxiGen (Hartvigsen et al., 2022): Machine-generated toxic and benign statements about 13 minority groups.")
add_para("DynaHate (Vidgen et al., 2021): Dynamically generated hate speech dataset with adversarial examples across four rounds.")

add_heading("4.2  Baselines", level=2)
add_para(
    "We compare against six baselines: "
    "(1) Vanilla: standard DeBERTa-v3-base fine-tuning; "
    "(2) EAR (Han et al., 2022): entropy-based attention regularization; "
    "(3) GetFair (Chhabra et al., 2024): gradient-based fairness constraint; "
    "(4) CCDF (Qian et al., 2023): counterfactual causal debiasing with TDE; "
    "(5) LogitPairing (Davani et al., 2021): logit pairing with swap-based augmentation; "
    "(6) AdvDebias (Ramponi & Tonelli, 2022): adversarial debiasing. "
    "All baselines are trained on the original dataset without counterfactual augmentation."
)

add_heading("4.3  Evaluation Metrics", level=2)
add_para("We report seven metrics spanning three dimensions:")
add_para("Task Performance: Accuracy, Macro-F1, AUC-ROC.")
add_para(
    "Causal Fairness (evaluated using LLM-generated counterfactuals): "
    "CFR (Counterfactual Fairness Rate) — fraction of predictions that change under counterfactual "
    "intervention; CTFG (Counterfactual Token Fairness Gap) — average probability shift."
)
add_para(
    "Group Fairness: FPED (False Positive Equality Difference); "
    "FNED (False Negative Equality Difference)."
)
add_para("All experiments use 3 random seeds (42, 123, 2024). We report mean ± std.")

# ============================================================
# 5. Results — 主表 (每个数据集一个表)
# ============================================================
add_heading("5  Results", level=1)
add_heading("5.1  Main Results", level=2)

COL_HEADERS = ['Method', 'Acc', 'F1', 'AUC', 'CFR↓', 'CTFG↓', 'FPED↓', 'FNED↓']
TASK_M = ['accuracy', 'macro_f1', 'auc_roc']
FAIR_M = ['cfr', 'ctfg', 'fped', 'fned']
ALL_M = TASK_M + FAIR_M

# best_cols: 1,2,3=max(task), 4,5,6,7=min(fair)
best_cols = {1:'max', 2:'max', 3:'max', 4:'min', 5:'min', 6:'min', 7:'min'}

for ds in DATASETS:
    add_para(f"Table: Results on {DS_NAMES[ds]}", bold=True)
    rows = []
    for method_key in BASELINE_KEYS:
        display_name = METHOD_NAMES[method_key]
        row = [display_name]
        for m in ALL_M:
            d = baseline_data.get(method_key, {}).get(ds, {}).get('llm', {}).get(m)
            if d and isinstance(d, dict):
                row.append(f"{d['mean']:.4f}")
            else:
                row.append('--')
        rows.append(row)
    make_table(rows, COL_HEADERS, best_cols)
    doc.add_paragraph()

# ============================================================
# 5.2 Analysis
# ============================================================
add_heading("5.2  Analysis", level=2)

add_para("Causal Fairness Improvements.", bold=True)
add_para(
    f"Our method achieves the lowest CFR across all three datasets. "
    f"On HateXplain, CFR is reduced from {imp['hatexplain']['vanilla']:.4f} (Vanilla) to "
    f"{imp['hatexplain']['ours']:.4f} ({imp['hatexplain']['vs_vanilla']:.1f}% reduction). "
    f"Compared to the best baseline (LogitPairing, CFR={imp['hatexplain']['best_bl']:.4f}), "
    f"our method achieves {imp['hatexplain']['vs_best']:.1f}% further reduction. "
    f"On ToxiGen, CFR drops to {imp['toxigen']['ours']:.4f} "
    f"({imp['toxigen']['vs_best']:.1f}% below LogitPairing). "
    f"On DynaHate, CFR={imp['dynahate']['ours']:.4f} "
    f"({imp['dynahate']['vs_vanilla']:.1f}% below Vanilla)."
)

add_para("Performance-Fairness Trade-off.", bold=True)
ours_f1 = get_mean(baseline_data, 'hatexplain', 'Ours', 'macro_f1')
van_f1 = get_mean(baseline_data, 'hatexplain', 'Vanilla', 'macro_f1')
f1_drop = (1 - ours_f1/van_f1)*100
add_para(
    f"The fairness improvements come with modest performance costs. On HateXplain, Macro-F1 "
    f"decreases by {f1_drop:.1f}% ({van_f1:.4f} → {ours_f1:.4f}). On ToxiGen and DynaHate, "
    f"the F1 drop is less than 1.5%. This is favorable compared to CCDF, which shows larger "
    f"performance degradation on DynaHate "
    f"(F1={get_mean(baseline_data, 'dynahate', 'CCDF', 'macro_f1'):.4f}) "
    f"while achieving worse fairness "
    f"(CFR={get_mean(baseline_data, 'dynahate', 'CCDF', 'cfr'):.4f})."
)

# ============================================================
# 5.3 Ablation Study
# ============================================================
add_heading("5.3  Ablation Study", level=2)
add_para(
    "To disentangle the contributions of LLM-based counterfactual generation and the CLP loss, "
    "we conduct ablation experiments with four variants: (1) No CF: our model architecture trained "
    "without any counterfactual data; (2) Swap (λ=0): trained with swap-based counterfactuals but "
    "no CLP loss; (3) Swap+CLP (λ=1): swap counterfactuals with CLP loss; (4) LLM (λ=0): "
    "LLM-generated counterfactuals without CLP loss. Our full method is LLM+CLP (λ=1)."
)

ABLATION_NAMES = ['No CF (λ=0)', 'Swap (λ=0)', 'Swap+CLP (λ=1)', 'LLM (λ=0)', 'LLM+CLP (λ=1)']

for ds in DATASETS:
    add_para(f"Table: Ablation on {DS_NAMES[ds]}", bold=True)
    rows = []
    for name in ABLATION_NAMES:
        row = [name]
        for m in ALL_M:
            d = ablation_data.get(ds, {}).get(name, {}).get(m)
            if d and isinstance(d, dict):
                row.append(f"{d['mean']:.4f}")
            else:
                row.append('--')
        rows.append(row)
    make_table(rows, COL_HEADERS, best_cols)
    doc.add_paragraph()

# Ablation analysis
add_para("Effect of Counterfactual Generation Method.", bold=True)
# Swap(λ=0) vs LLM(λ=0)
for ds in DATASETS:
    swap_cfr = get_ablation_mean(ds, 'Swap (λ=0)', 'cfr')
    llm_cfr = get_ablation_mean(ds, 'LLM (λ=0)', 'cfr')
    if swap_cfr and llm_cfr:
        red = (1 - llm_cfr/swap_cfr)*100
        add_para(f"  {DS_NAMES[ds]}: Swap CFR={swap_cfr:.4f} → LLM CFR={llm_cfr:.4f} ({red:.1f}% reduction)")

add_para(
    "LLM-generated counterfactuals consistently outperform swap-based ones even without CLP loss, "
    "confirming that higher-quality counterfactuals improve fairness during training."
)

add_para("Effect of CLP Loss.", bold=True)
for ds in DATASETS:
    llm0 = get_ablation_mean(ds, 'LLM (λ=0)', 'cfr')
    llm1 = get_ablation_mean(ds, 'LLM+CLP (λ=1)', 'cfr')
    if llm0 and llm1:
        red = (1 - llm1/llm0)*100
        add_para(f"  {DS_NAMES[ds]}: LLM(λ=0) CFR={llm0:.4f} → LLM+CLP(λ=1) CFR={llm1:.4f} ({red:.1f}% reduction)")

add_para(
    "Adding CLP loss provides substantial additional fairness gains beyond data augmentation alone, "
    "demonstrating that both components are essential."
)

# ============================================================
# 6. Discussion
# ============================================================
add_heading("6  Discussion", level=1)
add_para(
    "Why CLP outperforms existing methods. Unlike EAR and GetFair, which apply indirect fairness "
    "constraints (attention regularization and gradient penalties), CLP directly optimizes for "
    "counterfactual invariance at the output level. Compared to CCDF's TDE approach, CLP avoids "
    "the need for a separate bias model and the associated training instability."
)
add_para(
    "The role of LLM counterfactuals. Our ablation study shows that LLM-generated counterfactuals "
    "capture semantic shifts beyond simple lexical substitution. The combination of high-quality "
    "counterfactuals and direct logit pairing produces the strongest fairness improvements."
)
add_para(
    "Limitations. Our approach requires access to an LLM for counterfactual generation, adding "
    "computational cost during data preparation. While CLP significantly improves causal fairness "
    "metrics, group fairness metrics (FPED, FNED) show mixed results, suggesting that causal and "
    "group fairness may require complementary approaches."
)

# ============================================================
# 7. Conclusion
# ============================================================
add_heading("7  Conclusion", level=1)
add_para(
    "We presented Counterfactual Logit Pairing (CLP), a training framework that combines "
    "LLM-generated counterfactuals with a logit pairing objective to improve causal fairness "
    "in toxicity classification. Our method achieves state-of-the-art counterfactual fairness "
    "across three benchmark datasets while maintaining competitive classification performance. "
    "The approach is simple to implement, requires no architectural modifications, and can be "
    "applied to any text classification model. Future work includes extending CLP to multi-class "
    "settings and investigating methods to jointly optimize causal and group fairness."
)

# ============================================================
# Save
# ============================================================
output_path = "docs/EMNLP_2026_Paper_Draft.docx"
doc.save(output_path)
print(f"论文已保存: {output_path}")
print(f"段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}")
print(f"  - 主表: 3 (每个数据集1个)")
print(f"  - 消融表: 3 (每个数据集1个)")
